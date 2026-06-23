from __future__ import annotations
from lspopt import spectrogram_lspopt
from matplotlib.colors import Normalize
from matplotlib.backends.backend_pdf import PdfPages
import os
import sys
import subprocess
if sys.platform == 'win32':
    _original_popen = subprocess.Popen
    def _patched_popen(*args, **kwargs):
        # 0x08000000 - это системный флаг Windows CREATE_NO_WINDOW
        if 'creationflags' not in kwargs:
            kwargs['creationflags'] = 0x08000000
        return _original_popen(*args, **kwargs)
    subprocess.Popen = _patched_popen
from pathlib import Path
import re
import csv
import shutil
from typing import Optional, Dict, Tuple, Any
import traceback
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz  # PyMuPDF


def resource_path(relative_path):
    """ Получает абсолютный путь к ресурсу, работает для разработки и для PyInstaller """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)



RBD_ROOT = Path(resource_path(os.path.join("RBDtector", "RBDtector")))

if str(RBD_ROOT) not in sys.path:
    sys.path.insert(0, str(RBD_ROOT))
CAISR_ROOT = Path(resource_path("CAISR-App-main"))
if str(CAISR_ROOT) not in sys.path:
    sys.path.insert(0, str(CAISR_ROOT))

from app_logic.PSG_controller import PSGController
from util import settings

# if sys.stderr is None:
#     sys.stderr = open(os.devnull, "w", encoding="utf-8")
# if sys.stdout is None:
#     sys.stdout = open(os.devnull, "w", encoding="utf-8")

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

import mne
import yasa
import numpy as np
import pandas as pd
from PIL import Image

from sleepeegpy.pipeline import SpectralPipe
from mne.time_frequency import psd_array_multitaper

# ИМПОРТ НАШЕГО CAISR МОСТА
from caisr_bridge import run_full_caisr_pipeline, parse_caisr_nrem_events, plot_combined_hypnogram, plot_motor_events_standalone

import time as _time
import threading as _threading
import queue as _queue

# =====================================================================
# ПРОФИЛИРОВАНИЕ: Класс для замера времени этапов пайплайна
# =====================================================================
class PipelineTimer:
    """Lightweight stage timer that prints a summary report at the end."""
    def __init__(self):
        self.stages: list[tuple[str, float]] = []
        self._start: float = 0.0
        self._name: str = ""

    def start(self, name: str):
        self._start = _time.perf_counter()
        self._name = name
        print(f"[TIMER] Starting: {name}")

    def stop(self):
        elapsed = _time.perf_counter() - self._start
        self.stages.append((self._name, elapsed))
        print(f"[TIMER] {self._name}: {elapsed:.1f} sec")

    def report(self):
        if not self.stages:
            return
        total = sum(s[1] for s in self.stages)
        print(f"\n{'='*55}")
        print(f"PIPELINE TIMING REPORT (Total: {total:.1f} sec = {total/60:.1f} min)")
        print(f"{'='*55}")
        for name, elapsed in self.stages:
            pct = elapsed / max(total, 0.001) * 100
            bar = chr(9608) * int(pct / 2)
            print(f"  {name:40s} {elapsed:7.1f}s ({pct:5.1f}%) {bar}")
        print(f"{'='*55}\n")


def _convert_docx_to_pdf_safe(docx_path: Path, pdf_path: Path, timeout_sec: int = 120) -> bool:
    """Конвертация DOCX → PDF с таймаутом. Возвращает True при успехе."""
    result_q: _queue.Queue = _queue.Queue()

    def _worker():
        try:
            from docx2pdf import convert
            convert(str(docx_path), str(pdf_path))
            result_q.put(("ok", None))
        except Exception as e:
            result_q.put(("error", e))

    t = _threading.Thread(target=_worker, daemon=True)
    t.start()
    t.join(timeout=timeout_sec)

    if t.is_alive():
        print(f"WARNING: docx2pdf timed out after {timeout_sec}s — Word may be stuck. Skipping PDF conversion.")
        return False

    try:
        status, error = result_q.get_nowait()
    except _queue.Empty:
        return False

    if status == "error":
        print(f"WARNING: docx2pdf failed: {error}")
        return False
    return True


# Максимальное число страниц в Events Atlas PDF (защита от генерации сотен страниц matplotlib)
MAX_ATLAS_PAGES_PER_TYPE = 40

ADULT_CFG: Dict[str, Dict[str, Any]] = {
    "right": {
        "F": "E224", "C": "E183", "O": "E150",
        "EOG1": "E252", "EOG2": "E2",
        "EMG": "E240",
        "M": "E190",
        "M_NAME": "RM",
        "EEG_ALIAS": {"F": "EEG F4-RM", "C": "EEG C4-RM", "O": "EEG O2-RM"},
        "EOG1_ALIAS": "EOG E1-RM",
        "EOG2_ALIAS": "EOG E2-RM",
        "EMG_ALIAS": "EMG-R",
    },
    "left": {
        "F": "E36", "C": "E59", "O": "E116",
        "EOG1": "E226", "EOG2": "E47",
        "EMG": "E243",
        "M": "E94",
        "M_NAME": "LM",
        "EEG_ALIAS": {"F": "EEG F3-LM", "C": "EEG C3-LM", "O": "EEG O1-LM"},
        "EOG1_ALIAS": "EOG E1-LM",
        "EOG2_ALIAS": "EOG E2-LM",
        "EMG_ALIAS": "EMG-L",
    },
}

KIDS_CFG: Dict[str, Dict[str, Any]] = {
    "right": {
        "F": "E124", "C": "E104", "O": "E83",
        "EOG1": "E25", "EOG2": "E8",
        "EMG": "E126",
        "M": "E100",
        "M_NAME": "RM",
        "EEG_ALIAS": {"F": "EEG F4-RM", "C": "EEG C4-RM", "O": "EEG O2-RM"},
        "EOG1_ALIAS": "EOG E1-RM",
        "EOG2_ALIAS": "EOG E2-RM",
        "EMG_ALIAS": "EMG-R",
    },
    "left": {
        "F": "E24", "C": "E36", "O": "E70",
        "EOG1": "E226", "EOG2": "E47",
        "EMG": "E127",
        "M": "E58",
        "M_NAME": "LM",
        "EEG_ALIAS": {"F": "EEG F3-LM", "C": "EEG C3-LM", "O": "EEG O1-LM"},
        "EOG1_ALIAS": "EOG E1-LM",
        "EOG2_ALIAS": "EOG E2-LM",
        "EMG_ALIAS": "EMG-L",
    },
}

CFGS: Dict[str, Dict[str, Dict[str, Any]]] = {
    "adult_e256": ADULT_CFG,
    "kids_e128": KIDS_CFG,
}

_ENDTIME_RE = re.compile(r"(<endTime>)(\d+)(</endTime>)")
bad_pleth_times: list[list[int]] = []

_MALE_SLEEP_PARAM_RANGES = {
    "TST": ["390.5 ± 55.0", "375.0 ± 60.0", "337.7 ± 69.6", "330.0 ± 73.6", "323.1 ± 64.9", "314.2 ± 69.6", "331.6 ± 54.6", "330.9 ± 64.9", "308.1 ± 101.0", "289.6 ± 71.7"],
    "Eff": ["89.5 ± 7.0", "86.0 ± 9.0", "81.6 ± 11.4", "79.9 ± 12.9", "79.6 ± 10.4", "75.8 ± 13.78", "78.5 ± 9.78", "74.6 ± 11.68", "72.9 ± 20.18", "60.6 ± 13.5"],
    "Lat": ["14.5 ± 10.0", "15.0 ± 12.0", "13.9 ± 16.2", "17.5 ± 20.8", "16.8 ± 15.0", "14.6 ± 15.58", "14.8 ± 14.18", "21.5 ± 17.08", "34.5 ± 56.58", "49.0 ± 71.1"],
    "WASO": ["35.0 ± 20.0", "45.0 ± 25.0", "62.1 ± 41.9", "67.3 ± 49.7", "66.0 ± 35.2", "85.8 ± 47.68", "77.6 ± 38.68", "92.3 ± 40.88", "76.0 ± 30.28", "138.7 ± 38.4"],
    "S1%": ["4.5 ± 2.5", "4.8 ± 3.0", "5.0 ± 3.7", "5.7 ± 3.6", "6.1 ± 4.8", "5.9 ± 4.08", "4.6 ± 2.88", "5.2 ± 2.38", "4.2 ± 2.28", "8.1 ± 3.3"],
    "S2%": ["48.5 ± 8.0", "51.0 ± 9.0", "55.7 ± 11.1", "55.4 ± 9.7", "54.6 ± 9.1", "58.3 ± 11.78", "57.0 ± 7.48", "58.5 ± 8.98", "55.3 ± 17.78", "57.7 ± 13.9"],
    "SWS%": ["23.5 ± 6.5", "21.5 ± 7.0", "19.4 ± 8.9", "19.1 ± 7.5", "20.1 ± 7.6", "18.6 ± 9.48", "19.9 ± 7.58", "17.3 ± 7.88", "23.1 ± 16.08", "16.1 ± 8.7"],
    "REM%": ["23.5 ± 5.0", "22.0 ± 6.0", "19.9 ± 6.8", "19.8 ± 6.5", "19.1 ± 5.9", "17.2 ± 8.08", "18.4 ± 5.98", "19.0 ± 4.98", "17.5 ± 6.98", "18.1 ± 6.3"],
    "B90": ["1.0 ± 3.0", "3.5 ± 10.0", "10.3 ± 35.3", "11.3 ± 33.7", "5.3 ± 11.3", "15.4 ± 31.48", "13.4 ± 20.98", "14.8 ± 33.88", "54.1 ± 86.28", "11.0 ± 12.6"],
    "AvS": ["96.5 ± 1.0", "95.5 ± 1.5", "94.3 ± 2.8", "94.5 ± 1.8", "95.0 ± 1.5", "93.5 ± 1.78", "93.9 ± 1.48", "94.0 ± 1.98", "92.8 ± 2.68", "92.8 ± 1.4"],
    "MinS": ["92.0 ± 3.5", "90.0 ± 4.5", "86.3 ± 7.4", "86.5 ± 5.3", "87.1 ± 5.4", "83.6 ± 6.88", "84.5 ± 6.08", "83.9 ± 8.98", "83.9 ± 4.68", "84.1 ± 2.5"],
}

_FEMALE_SLEEP_PARAM_RANGES = {
    "TST": ["405.0 ± 50.0", "385.0 ± 65.0", "338.2 ± 84.6", "343.2 ± 75.8", "312.1 ± 72.9", "333.1 ± 71.5", "327.4 ± 49.7", "305.6 ± 87.6", "281.9 ± 61.2", "291.7 ± 65.5"],
    "Eff": ["91.0 ± 6.0", "88.5 ± 8.5", "83.3 ± 12.8", "82.2 ± 12.3", "76.9 ± 14.6", "78.8 ± 10.4", "79.3 ± 9.6", "75.0 ± 16.1", "72.7 ± 16.0", "67.8 ± 12.4"],
    "Lat": ["12.5 ± 9.0", "13.0 ± 10.0", "13.1 ± 11.9", "17.3 ± 21.3", "18.6 ± 21.3", "19.0 ± 19.8", "19.1 ± 15.6", "20.0 ± 33.6", "25.0 ± 29.3", "27.5 ± 26.6"],
    "WASO": ["30.0 ± 18.0", "40.0 ± 22.0", "55.0 ± 48.7", "57.1 ± 42.2", "77.1 ± 55.3", "73.3 ± 43.0", "68.2 ± 36.7", "79.7 ± 40.3", "84.4 ± 58.8", "113.1 ± 53.4"],
    "S1%": ["3.8 ± 2.0", "4.0 ± 2.2", "4.1 ± 2.3", "4.3 ± 2.7", "4.6 ± 3.2", "5.1 ± 4.0", "5.1 ± 3.0", "4.8 ± 3.8", "5.1 ± 2.1", "5.6 ± 3.1"],
    "S2%": ["47.0 ± 7.0", "50.0 ± 8.0", "54.4 ± 7.7", "53.3 ± 9.5", "54.5 ± 9.5", "56.4 ± 11.7", "55.1 ± 9.8", "55.8 ± 12.3", "55.5 ± 8.2", "58.5 ± 13.1"],
    "SWS%": ["25.5 ± 6.0", "24.0 ± 6.5", "22.1 ± 6.8", "22.7 ± 9.2", "24.3 ± 7.6", "20.3 ± 9.8", "22.0 ± 7.4", "23.5 ± 10.7", "23.4 ± 8.8", "19.4 ± 11.3"],
    "REM%": ["23.7 ± 5.0", "22.0 ± 5.5", "19.3 ± 6.2", "19.7 ± 7.0", "16.7 ± 7.1", "18.2 ± 6.8", "17.9 ± 5.6", "15.9 ± 6.0", "16.0 ± 7.4", "16.5 ± 6.4"],
    "B90": ["0.2 ± 1.0", "0.4 ± 1.2", "0.5 ± 1.4", "4.0 ± 16.5", "6.3 ± 15.8", "18.7 ± 56.7", "18.1 ± 39.8", "16.7 ± 26.7", "27.9 ± 61.3", "26.2 ± 43.6"],
    "AvS": ["97.0 ± 0.8", "96.5 ± 1.0", "96.0 ± 1.2", "95.4 ± 1.6", "94.3 ± 1.6", "94.2 ± 2.3", "94.0 ± 1.9", "93.3 ± 2.0", "93.0 ± 1.9", "93.1 ± 1.9"],
    "MinS": ["93.5 ± 2.5", "92.0 ± 3.5", "90.7 ± 4.7", "88.1 ± 5.4", "86.3 ± 5.2", "85.5 ± 6.3", "84.1 ± 5.3", "83.4 ± 5.5", "82.1 ± 6.0", "83.9 ± 8.1"],
}


def _choose_start_dt(raw: mne.io.BaseRaw) -> pd.Timestamp:
    meas_date = raw.info.get("meas_date", None)
    if meas_date is None:
        return pd.Timestamp.now().normalize() + pd.Timedelta(hours=22)
    return pd.Timestamp(meas_date)

def _ensure_loaded(raw: mne.io.BaseRaw) -> mne.io.BaseRaw:
    if not raw.preload:
        raw.load_data()
    return raw

def _save_kv_csv(path: Path, mapping: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerows(mapping.items())

def _write_sleep_profile_txt(path: Path, *, start_dt: pd.Timestamp, hypno_1hz_int: np.ndarray) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    hypno_1hz_int = np.asarray(hypno_1hz_int, dtype=int)

    if hypno_1hz_int.ndim != 1:
        hypno_1hz_int = hypno_1hz_int.ravel()

    if hypno_1hz_int.size == 0:
        with open(path, "w", encoding="utf-8") as f:
            f.write(f"Start Time: {start_dt.strftime('%d.%m.%Y %H:%M:%S')}\n")
            f.write("Generated-By: yasa-headless\n")
            f.write("\n")
        return

    epoch_sec = 30
    valid_stage_order = [0, 1, 2, 3, 4]

    def aggregate_epoch(epoch_vals: np.ndarray) -> int:
        epoch_vals = np.asarray(epoch_vals, dtype=int)
        epoch_vals = epoch_vals[np.isin(epoch_vals, valid_stage_order)]

        if epoch_vals.size == 0:
            return -1

        counts = np.bincount(epoch_vals, minlength=max(valid_stage_order) + 1)
        max_count = counts.max()
        winners = np.flatnonzero(counts == max_count)

        for st in (4, 3, 2, 1, 0):
            if st in winners:
                return int(st)

        return int(winners[0])

    n_full_epochs = hypno_1hz_int.size // epoch_sec
    remainder = hypno_1hz_int.size % epoch_sec

    hypno_int_30s = []
    if n_full_epochs > 0:
        full = hypno_1hz_int[: n_full_epochs * epoch_sec].reshape(n_full_epochs, epoch_sec)
        hypno_int_30s.extend(aggregate_epoch(epoch) for epoch in full)

    if remainder > 0:
        tail = hypno_1hz_int[n_full_epochs * epoch_sec :]
        hypno_int_30s.append(aggregate_epoch(tail))

    hypno_int_30s = np.asarray(hypno_int_30s, dtype=int)

    def int_stage_to_str(v: int) -> str:
        mapping = {-2: "UNS", -1: "ART", 0: "W", 1: "N1", 2: "N2", 3: "N3", 4: "REM"}
        return mapping.get(int(v), "ART")

    hypno_str_30s = [int_stage_to_str(v) for v in hypno_int_30s]
    hypno_str_1hz = np.repeat(hypno_str_30s, epoch_sec)[: hypno_1hz_int.size]

    with open(path, "w", encoding="utf-8") as f:
        f.write(f"Start Time: {start_dt.strftime('%d.%m.%Y %H:%M:%S')}\n")
        f.write("Generated-By: yasa-headless\n")
        f.write("\n")
        for i, st in enumerate(hypno_str_1hz):
            t = (start_dt + pd.Timedelta(seconds=int(i))).strftime("%H:%M:%S,%f")[:-3]
            f.write(f"{t}; {st}\n")


def _write_minimal_event_file(path: Path, *, start_dt: pd.Timestamp, label: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    t = start_dt.strftime("%H:%M:%S,%f")[:-3]
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"Start Time: {start_dt.strftime('%d.%m.%Y %H:%M:%S')}\n")
        f.write("Generated-By: headless\n")
        f.write("\n")
        f.write(f"{t} - {t}; 0; {label}\n")

def fix_mff_epochs_xml(mff_path: Path) -> None:
    """Recover or recreate a corrupted/missing epochs.xml for EGI .mff files.

    Parses signal1.bin via MNE's internal ``_get_blocks`` to obtain the exact
    block count and sample total, then reads the time-divisor from ``info.xml``
    to compute the precise ``endTime`` in microseconds.

    Parameters
    ----------
    mff_path : Path
        Path to the .mff directory.
    """
    import xml.etree.ElementTree as ET
    from mne.io.egi.general import _get_blocks

    mff_path = Path(mff_path)
    signal_bin   = mff_path / "signal1.bin"
    epochs_xml   = mff_path / "epochs.xml"
    info_xml_path = mff_path / "info.xml"

    if not signal_bin.exists() or not info_xml_path.exists():
        raise FileNotFoundError(
            f"Missing essential MFF files (signal1.bin / info.xml) in {mff_path}"
        )

    # 1. Use MNE's built-in parser — guarantees perfect alignment with MNE's expectations
    print(f"Scanning blocks to calculate exact size for {mff_path.name}...")
    signal_blocks  = _get_blocks(str(signal_bin))
    total_samples  = int(signal_blocks["samples_block"].sum())
    last_block_idx = int(signal_blocks["n_blocks"])   # exact, not hardcoded 9999
    sfreq          = int(signal_blocks["sfreq"])

    # 2. Extract time divisor from info.xml
    root = ET.parse(info_xml_path).getroot()
    record_time_elem = root.find(".//{http://www.egi.com/info_mff}recordTime")
    if record_time_elem is None:
        raise ValueError("Could not find <recordTime> in info.xml")

    record_time = record_time_elem.text
    match = re.match(r".*\.(\d{6}(?:\d{3})?)[+-]", record_time)
    if not match:
        raise ValueError(f"Unexpected recordTime format: {record_time}")

    frac = match.group(1)
    div  = 1000 if len(frac) == 6 else 1000000   # microseconds vs nanoseconds

    # 3. Calculate exact endTime in microseconds
    duration_us = int((total_samples / sfreq) * (div * 1000))

    # 4. Backup existing corrupted epochs.xml (if present)
    if epochs_xml.exists():
        backup_path = mff_path.parent / f"{mff_path.name}__epochs_corrupted.xml"
        if not backup_path.exists():
            shutil.copy2(epochs_xml, backup_path)

    # 5. Write corrected epochs.xml
    xml_str = (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<epochs xmlns="http://www.egi.com/epochs_mff">\n'
        '    <epoch>\n'
        f'        <beginTime>0</beginTime>\n'
        f'        <endTime>{duration_us}</endTime>\n'
        f'        <firstBlock>1</firstBlock>\n'
        f'        <lastBlock>{last_block_idx}</lastBlock>\n'
        '    </epoch>\n'
        '</epochs>\n'
    )
    epochs_xml.write_text(xml_str, encoding="utf-8")
    print(
        f"Successfully created/recovered epochs.xml for {mff_path.name} "
        f"with {total_samples} samples and {last_block_idx} blocks "
        f"(endTime={duration_us} µs)."
    )


def safe_read_mff(
    mff_path: str | Path,
    preload: bool = False,
    verbose: bool = False,
    step: int = 1000,
    max_k: int = 20000,
    **kwargs,
) -> mne.io.BaseRaw:
    """Read an EGI .mff file, auto-recovering a corrupted/missing epochs.xml.

    On first failure the function calls :func:`fix_mff_epochs_xml` to
    rebuild ``epochs.xml`` from the binary signal data, then retries.
    """
    mff_path   = Path(mff_path)
    epochs_xml = mff_path / "epochs.xml"

    try:
        return mne.io.read_raw_egi(str(mff_path), preload=preload, verbose=verbose, **kwargs)
    except Exception as first_err:
        if not epochs_xml.exists():
            print(f"[Recovery] epochs.xml missing for {mff_path.name} — rebuilding...")
        else:
            print(f"[Recovery] Initial read failed for {mff_path.name} — rebuilding epochs.xml...")
            if verbose:
                print(f"  Original error: {first_err}")

        fix_mff_epochs_xml(mff_path)
        return mne.io.read_raw_egi(str(mff_path), preload=preload, verbose=verbose, **kwargs)



def _save_hypnogram_png(path: Path, hypno_1hz_int: np.ndarray, title: str | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    hypno_1hz_int = np.asarray(hypno_1hz_int, dtype=int)
    t_hours = np.arange(hypno_1hz_int.size) / 3600.0

    remapped = pd.Series(hypno_1hz_int).map({-2: -2, -1: -1, 0: 0, 1: 2, 2: 3, 3: 4, 4: 1}).fillna(-2).to_numpy()
    y = -1.0 * remapped

    fig, ax = plt.subplots(figsize=(14, 4), dpi=150)
    ax.step(t_hours, y, where="post", color="k", linewidth=1.8)

    masks = {
        "REM": np.ma.masked_not_equal(remapped, 1),
        "W": np.ma.masked_not_equal(remapped, 0),
        "N1": np.ma.masked_not_equal(remapped, 2),
        "N2": np.ma.masked_not_equal(remapped, 3),
        "N3": np.ma.masked_not_equal(remapped, 4),
    }
    colors = {"REM": "r", "W": "b", "N1": "c", "N2": "g", "N3": "m"}

    for stage, arr in masks.items():
        ax.step(t_hours, -1.0 * arr, where="post", color=colors[stage], linewidth=2.2)

    ax.set_yticks([0, -1, -2, -3, -4])
    ax.set_yticklabels(["W", "R", "N1", "N2", "N3"])
    ax.set_ylim(-4.5, 0.5)
    ax.set_xlim(0, max(float(t_hours[-1]) if t_hours.size else 0.0, 0.01))
    ax.set_xlabel("Time [hrs]")
    if title:
        ax.set_title(title)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def _harmonic_notch_freqs(*, sfreq: float, base_freq: float = 50.0, max_freq: float | None = None) -> np.ndarray:
    nyquist = float(sfreq) / 2.0
    upper = nyquist - 1e-6 if max_freq is None else min(float(max_freq), nyquist - 1e-6)
    if base_freq <= 0 or upper < base_freq:
        return np.array([], dtype=float)
    n_max = int(np.floor(upper / float(base_freq)))
    if n_max < 1:
        return np.array([], dtype=float)
    return np.arange(1, n_max + 1, dtype=float) * float(base_freq)


def _apply_mff_prefilter(
    raw: mne.io.BaseRaw,
    *,
    picks: list[str],
    l_freq: float = 0.1,
    h_freq: float = 70.0,
    notch_base: float = 50.0,
    verbose: bool = True,
) -> mne.io.BaseRaw:
    raw = raw.copy()
    _ensure_loaded(raw)

    picks_existing = [ch for ch in picks if ch in raw.ch_names]
    if not picks_existing:
        return raw

    sfreq = float(raw.info["sfreq"])
    notch_freqs = _harmonic_notch_freqs(sfreq=sfreq, base_freq=notch_base)
    notch_freqs = notch_freqs[notch_freqs < (sfreq / 2.0)]

    if verbose:
        notch_msg = ", ".join(f"{x:.0f}" for x in notch_freqs) if notch_freqs.size else "none"
        print(f"Applying prefilter to {len(picks_existing)} channels: bandpass {l_freq}-{h_freq} Hz; notch {notch_msg} Hz")

    if notch_freqs.size:
        raw.notch_filter(freqs=notch_freqs, picks=picks_existing, n_jobs=1, verbose=False)

    raw.filter(l_freq=l_freq, h_freq=h_freq, picks=picks_existing, n_jobs=1, verbose=False)
    return raw


def hypno_report_stats(hypno_1hz_int: np.ndarray) -> dict[str, float]:
    hypno = list(np.asarray(hypno_1hz_int, dtype=int))
    data_length = len(hypno)
    wake = hypno.count(0)
    n1 = hypno.count(1)
    n2 = hypno.count(2)
    n3 = hypno.count(3)
    rem = hypno.count(4)
    art = hypno.count(-1)

    sleep_time = n1 + n2 + n3 + rem + art
    sleep_time_wo_art = max(n1 + n2 + n3 + rem, 1)

    first_time_n1 = hypno.index(1) if 1 in hypno else 0
    data_after_1sleep = hypno[first_time_n1:]
    wake_after_1sleep = data_after_1sleep.count(0)
    wake_after_1sleep_percent = wake_after_1sleep / max(data_length, 1) * 100.0
    sleep_efficiency = sleep_time / max(data_length, 1) * 100.0




    return {
        "TIB": float(data_length),
        "SPT": float(sleep_time),
        "LAT": float(first_time_n1),
        "WASO": float(wake_after_1sleep),
        "WASO%": float(wake_after_1sleep_percent),
        "SE": float(sleep_efficiency),
        "Wake": float(wake / sleep_time_wo_art * 100.0),
        "N1": float(n1 / sleep_time_wo_art * 100.0),
        "N2": float(n2 / sleep_time_wo_art * 100.0),
        "N3": float(n3 / sleep_time_wo_art * 100.0),
        "REM": float(rem / sleep_time_wo_art * 100.0),
    }


def plot_sleep_stages(subject_id: str, hypno_dict: dict[str, float], out_path: Path) -> Path:
    fig, ax = plt.subplots(figsize=(6, 6))
    labels = ["Wake", "N1", "N2", "N3", "REM"]
    values = [hypno_dict["Wake"], hypno_dict["N1"], hypno_dict["N2"], hypno_dict["N3"], hypno_dict["REM"]]
    ax.pie(
        values,
        labels=labels,
        autopct="%.1f%%",
        wedgeprops={"linewidth": 3.0, "edgecolor": "white"},
        textprops={"size": "x-large"},
    )
    ax.set_title(f"Sleep Stages - {subject_id}", fontsize=18)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, bbox_inches="tight", dpi=150)
    plt.close(fig)
    return out_path


def get_result_ranges(age: int | float, sex: str) -> dict[str, str]:
    sleep_stages_percentage = _FEMALE_SLEEP_PARAM_RANGES if str(sex).upper() == "F" else _MALE_SLEEP_PARAM_RANGES

    if age < 30:
        i = 0
    elif age < 40:
        i = 1
    elif age <= 44:
        i = 2
    elif age <= 49:
        i = 3
    elif age <= 54:
        i = 4
    elif age <= 59:
        i = 5
    elif age <= 64:
        i = 6
    elif age <= 69:
        i = 7
    elif age <= 74:
        i = 8
    else:
        i = 9

    return {
        "N1": sleep_stages_percentage["S1%"][i],
        "N2": sleep_stages_percentage["S2%"][i],
        "N3": sleep_stages_percentage["SWS%"][i],
        "REM": sleep_stages_percentage["REM%"][i],
        "TST": sleep_stages_percentage["TST"][i],
        "LAT": sleep_stages_percentage["Lat"][i],
        "WASO": sleep_stages_percentage["WASO"][i],
        "SE": sleep_stages_percentage["Eff"][i],
        "B90": sleep_stages_percentage["B90"][i],
        "AvS": sleep_stages_percentage["AvS"][i],
        "MinS": sleep_stages_percentage["MinS"][i],
    }


def _seconds_to_hms_str(seconds: float | int) -> str:
    seconds = int(max(0, round(float(seconds))))
    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60
    return f"{h:02d}:{m:02d}:{s:02d}"


def get_local_min_max(sec: int, spo_pleth: np.ndarray, time_span: int):
    from scipy.signal import find_peaks

    num_points = int(sec) * 250
    pleth_span = spo_pleth[num_points:(num_points + time_span * 250)]
    local_max = find_peaks(pleth_span, distance=175)[0]
    local_min = find_peaks(-1 * pleth_span, distance=175)[0]

    df_pleth = pd.DataFrame(pleth_span)
    df_local_min = df_pleth.iloc[local_min, :].copy()
    df_local_max = df_pleth.iloc[local_max, :].copy()
    df_local_min.rename(columns={0: "value"}, inplace=True)
    df_local_max.rename(columns={0: "value"}, inplace=True)
    return df_local_min, df_local_max


def is_pleth_wrong_amplitude(time: int, spo_pleth: np.ndarray, time_interval: int) -> bool:
    df_min, df_max = get_local_min_max(time, spo_pleth, time_interval)
    if len(df_min) < time_interval * 45 / 60 or len(df_max) < time_interval * 45 / 60:
        return True

    sensitivity_std = 0.45
    min_ratio = abs(df_min.std().value / df_min.mean().value) if df_min.mean().value != 0 else np.inf
    max_ratio = abs(df_max.std().value / df_max.mean().value) if df_max.mean().value != 0 else np.inf

    if min_ratio > sensitivity_std or max_ratio > sensitivity_std:
        if min_ratio > sensitivity_std and len(df_min) > 0:
            if df_min["value"].iloc[0] > df_min.mean().value + 2 * df_min.std().value:
                df_min = df_min.drop(df_min.index[0])
            if len(df_min) > 0 and df_min["value"].iloc[-1] > df_min.mean().value + 2 * df_min.std().value:
                df_min = df_min.drop(df_min.index[-1])

        if max_ratio > sensitivity_std and len(df_max) > 0:
            if df_max["value"].iloc[0] < df_max.mean().value - 2 * df_max.std().value:
                df_max = df_max.drop(df_max.index[0])
            if len(df_max) > 0 and df_max["value"].iloc[-1] < df_max.mean().value - 2 * df_max.std().value:
                df_max = df_max.drop(df_max.index[-1])

        min_ratio = abs(df_min.std().value / df_min.mean().value) if len(df_min) and df_min.mean().value != 0 else np.inf
        max_ratio = abs(df_max.std().value / df_max.mean().value) if len(df_max) and df_max.mean().value != 0 else np.inf
        if min_ratio > sensitivity_std or max_ratio > sensitivity_std:
            return True

    df_min = df_min.copy()
    df_max = df_max.copy()
    df_min["type"] = "min"
    df_max["type"] = "max"
    concated_df = pd.concat([df_max, df_min]).sort_index().reset_index(level=0)
    len_concated_df = len(concated_df)
    if len_concated_df == 0:
        return True

    count_wrong_events = 0.0
    sensitivity_dif = 100
    sensitivity_max = 100
    sensitivity_min = -100
    sensitivity_result = 0.12

    for index, row in concated_df.iterrows():
        if (row["type"] == "max" and row["value"] < sensitivity_max) or (row["type"] == "min" and row["value"] > sensitivity_min):
            count_wrong_events += 1.0
        elif index <= len_concated_df - 2:
            next_index = index + 1
            if concated_df.iloc[index]["type"] == concated_df.iloc[next_index]["type"] and next_index < len_concated_df - 1:
                next_index += 1
                if concated_df.iloc[index]["type"] != concated_df.iloc[next_index]["type"]:
                    difference_val = abs(abs(concated_df.iloc[index]["value"]) - abs(concated_df.iloc[next_index]["value"]))
                    if difference_val < sensitivity_dif:
                        count_wrong_events += 1.0

    return sensitivity_result < count_wrong_events / len_concated_df


def is_pleth_wrong_fft(time: int, spo_pleth: np.ndarray, time_interval: int) -> bool:
    from scipy import fftpack

    sensitivity_percentage_of_maximum = 0.25
    yf = fftpack.fft(spo_pleth[time * 250:(time + time_interval) * 250])
    return bool(
        np.max(np.abs(yf[:8])) > sensitivity_percentage_of_maximum * np.max(np.abs(yf[8:17]))
        or np.max(np.abs(yf[17:20])) > sensitivity_percentage_of_maximum * np.max(np.abs(yf[8:17]))
    )


def is_pleth_wrong(time: int, spo_pleth: np.ndarray, time_interval: int) -> bool:
    wrong = is_pleth_wrong_amplitude(time, spo_pleth, time_interval)
    if not wrong:
        wrong = is_pleth_wrong_fft(time, spo_pleth, time_interval)
    return wrong


def _create_pleth_intervals(spo_pleth: np.ndarray) -> tuple[list[list[int]], list[list[int]]]:
    global bad_pleth_times
    good_pleth_times: list[list[int]] = []
    bad_pleth_times = []

    start_time = 300
    time_interval = 10
    interval_list = np.arange(start_time, int(len(spo_pleth[0]) / 250) - time_interval, time_interval)

    for time in interval_list:
        wrong = is_pleth_wrong(int(time), spo_pleth[0], time_interval)
        if wrong:
            bad_pleth_times.append([int(time), int(time + time_interval)])
        else:
            good_pleth_times.append([int(time), int(time + time_interval)])
    return bad_pleth_times, good_pleth_times


def analyse_spo2(path: str | Path, spo2_channel: str = "SaO2 SPO2", pleth_channel: str = "Pulse Pleth"):
    path = str(path)
    sum_spo2 = 0.0
    events_avg_low = []
    minimal_spo2 = [0, 100.0]
    minimal_spo2_avg = [0, 100.0]
    num_events_under_90 = 0
    events_under_90 = []

    edf = mne.io.read_raw_edf(input_fname=path, preload=False, verbose="ERROR")
    spo_pleth = edf[pleth_channel][0] * 1000000
    good_pleth_times = _create_pleth_intervals(spo_pleth)[1]
    spo_values = edf[spo2_channel][0] * 1000000

    extra_stats = [events_avg_low, minimal_spo2_avg, events_under_90]
    for time_span in good_pleth_times:
        spo_info = spo_values[0][time_span[0] * 250:time_span[1] * 250]
        if spo_info.size == 0:
            continue

        sum_spo2 += float(np.sum(spo_info))
        if minimal_spo2[1] > float(np.min(spo_info)):
            minimal_spo2 = [time_span[0], float(np.min(spo_info))]

        below_90 = spo_info < 90
        num_events_under_90 += int(np.sum(below_90))
        if np.any(below_90):
            for val in spo_info[below_90]:
                events_under_90.append([time_span, float(val)])

        avg = float(np.average(spo_info))
        if avg < 93:
            events_avg_low.append([time_span[0], avg])
        if minimal_spo2_avg[1] > avg:
            minimal_spo2_avg = [time_span[0], avg]
        extra_stats = [events_avg_low, minimal_spo2_avg, events_under_90]

    denom = max((len(good_pleth_times) * 10) * 250, 1)
    avg_spo2 = sum_spo2 / denom
    time_under_90 = num_events_under_90 * 0.004
    return avg_spo2, minimal_spo2, time_under_90, extra_stats

def _check_if_abnormal(val_str: str, normal_str: str) -> bool:
    """
    Парсит строки вида 'Mean ± SD' и проверяет, выходит ли значение за пределы 2 SD.
    Возвращает True, если значение аномально, иначе False.
    """
    try:
        # Извлекаем числовое значение из результата (например, "85.4%" -> 85.4)
        val_match = re.search(r"[-+]?\d*\.\d+|\d+", str(val_str))
        if not val_match:
            return False
        val = float(val_match.group())

        # Извлекаем среднее и SD из референсной строки (например, "338.2 ± 84.6")
        norm_match = re.search(r"([-+]?\d*\.\d+|\d+)\s*±\s*([-+]?\d*\.\d+|\d+)", str(normal_str))
        if norm_match:
            mean = float(norm_match.group(1))
            sd = float(norm_match.group(2))
            # Проверка на выход за пределы 2 стандартных отклонений (~95% доверительный интервал)
            if val < (mean - 2 * sd) or val > (mean + 2 * sd):
                return True
        return False
    except Exception:
        return False

import re
import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _check_if_abnormal(val_str: str, normal_str: str) -> bool:
    """
    Парсит строки вида 'Mean ± SD' и проверяет, выходит ли значение за пределы 2 SD.
    Возвращает True, если значение аномально, иначе False.
    """
    try:
        val_match = re.search(r"[-+]?\d*\.\d+|\d+", str(val_str))
        if not val_match:
            return False
        val = float(val_match.group())

        norm_match = re.search(r"([-+]?\d*\.\d+|\d+)\s*±\s*([-+]?\d*\.\d+|\d+)", str(normal_str))
        if norm_match:
            mean = float(norm_match.group(1))
            sd = float(norm_match.group(2))
            if val < (mean - 2 * sd) or val > (mean + 2 * sd):
                return True
        return False
    except Exception:
        return False
def calculate_plm_metrics(caisr_nrem_events: pd.DataFrame, hypno_1hz_int: np.ndarray) -> dict:
    """
    Calculate Limb Movement parameters according to AASM standards.
    A valid PLM sequence requires >= 4 consecutive movements, separated by 5 to 90 seconds.
    """
    # Calculate Total Sleep Time (TST) in hours (stages 1, 2, 3, 4)
    tst_sec = int(np.sum(np.isin(hypno_1hz_int, [1, 2, 3, 4])))
    tst_hours = tst_sec / 3600.0

    if caisr_nrem_events is None or caisr_nrem_events.empty or tst_hours <= 0:
        return {"PLMI": 0.0, "Total_PLM": 0, "LMSI": 0.0, "Total_LMS": 0}

    events = caisr_nrem_events.sort_values(by="start_sec").reset_index(drop=True)
    total_lms = len(events)
    
    plm_count = 0
    current_sequence = 1
    
    # Iterate through events to find sequences adhering to AASM 5-90s rule
    for i in range(1, total_lms):
        interval = events.loc[i, "start_sec"] - events.loc[i-1, "start_sec"]
        
        if 5.0 <= interval <= 90.0:
            current_sequence += 1
        else:
            # If sequence breaks, check if it met the >= 4 criterion
            if current_sequence >= 4:
                plm_count += current_sequence
            current_sequence = 1
            
    # Check the final sequence
    if current_sequence >= 4:
        plm_count += current_sequence

    return {
        "PLMI": plm_count / tst_hours,
        "Total_PLM": plm_count,
        "LMSI": total_lms / tst_hours,  # Raw Limb Movement Index (not periodic)
        "Total_LMS": total_lms
    }


import fitz  # PyMuPDF

import fitz  # PyMuPDF
from pathlib import Path

def extract_watchpat_visuals(pdf_path: Path, out_dir: Path) -> list[Path]:
    """
    Извлекает визуалы WatchPAT с финальной подстройкой границ:
    - Стр 2: Шкала pAHI
    - Стр 3: Блок графиков (еще -30 единиц сверху)
    """
    extracted_paths = []
    try:
        doc = fitz.open(str(pdf_path))
        
        # --- 1. ИЗВЛЕЧЕНИЕ ШКАЛЫ pAHI (СТРАНИЦА 2) ---
        page2 = doc[1]
        text_instances_pahi = page2.search_for("pAHI")
        if text_instances_pahi:
            inst = sorted(text_instances_pahi, key=lambda x: x.y1)[-1]
            crop_pahi = fitz.Rect(30, inst.y0 - 10, 580, inst.y1 + 110)
            pix2 = page2.get_pixmap(clip=crop_pahi, dpi=600)
            pahi_path = out_dir / f"pAHI_scale_{pdf_path.stem}.png"
            pix2.save(str(pahi_path))
            extracted_paths.append(pahi_path)

        # --- 2. ИЗВЛЕЧЕНИЕ ГРАФИКОВ (СТРАНИЦА 3) С ДОПОЛНИТЕЛЬНОЙ ОБРЕЗКОЙ СВЕРХУ ---
        if len(doc) > 2:
            page3 = doc[2]
            top_anchor = page3.search_for("PAT Respiratory Events")
            bottom_anchor = page3.search_for("Wake / Sleep stages")
            
            if top_anchor and bottom_anchor:
                # ПРИМЕНЯЕМ ТВОИ КОРРЕКТИРОВКИ:
                # y0: был -15 -> стал +15 (сместили верхнюю границу вниз на 30 единиц)
                # x0 и x1 оставляем расширенными (15 и 595)
                crop_rect = fitz.Rect(
                    15,                    # Слева (широко)
                    top_anchor[0].y0 + 55, # Сверху (срезали еще 30 пикселей вниз)
                    595,                   # Справа (широко)
                    bottom_anchor[0].y1 + 160 # Низ (полная гипнограмма)
                )
            else:
                crop_rect = fitz.Rect(15, 95, 595, 780)

            pix3 = page3.get_pixmap(clip=crop_rect, dpi=600)
            graphs_path = out_dir / f"full_respiratory_graphs_{pdf_path.stem}.png"
            pix3.save(str(graphs_path))
            extracted_paths.append(graphs_path)

        doc.close()
    except Exception as e:
        print(f"Ошибка при извлечении: {e}")
    
    return extracted_paths

def parse_watchpat_pdf(pdf_path: "str | Path") -> dict:
    doc = fitz.open(str(pdf_path))
    raw_text = ""
    for page in doc:
        blocks = page.get_text("blocks")
        blocks.sort(key=lambda b: (b[1], b[0]))
        for b in blocks:
            raw_text += b[4] + " "
    doc.close()

    clean_text = re.sub(r'\s+', ' ', raw_text)

    data = {
        "date": "N/A",
        "trt_hours": 0.0,
        "tst_hours": 0.0,
        "efficiency": 0.0,
        "wakes_per_hour": 0.0,
        "pAHI": 0.0,
        "ODI": 0.0,
        "mean_sat": 0.0,
        "min_sat": 0.0,
        "sat_below_90_pct": 0.0,
        "snoring_mean_db": 0.0,
        "hr_stats": None,
        "pos_stats": None,
        "resp_stats": None,
    }

    m = re.search(r"(\d{2}/\d{2}/\d{4})", clean_text)
    if m:
        data["date"] = m.group(1)

    m = re.search(r"Total Recording Time[^\d]+(\d+)\s*hrs[,\s]+(\d+)\s*min", clean_text, re.IGNORECASE)
    if m:
        data["trt_hours"] = int(m.group(1)) + int(m.group(2)) / 60.0

    m = re.search(r"Total Sleep Time[^\d]+(\d+)\s*hrs[,\s]+(\d+)\s*min", clean_text, re.IGNORECASE)
    if m:
        data["tst_hours"] = int(m.group(1)) + int(m.group(2)) / 60.0

    if data["trt_hours"] > 0:
        data["efficiency"] = (data["tst_hours"] / data["trt_hours"]) * 100.0

    m = re.search(r"Number of Wakes\s*:?\s*(\d+)", clean_text, re.IGNORECASE)
    if m and data["tst_hours"] > 0:
        data["wakes_per_hour"] = int(m.group(1)) / data["tst_hours"]

    m = re.search(r"pAHI\s*=\s*(\d+(?:\.\d+)?)", clean_text, re.IGNORECASE)
    if m:
        data["pAHI"] = float(m.group(1))
    else:
        m = re.search(
            r"pAHI\s*3%\s*:?\s*\d+\s+(?:\d+\.\d+\s+)*(\d+\.\d+)",
            clean_text, re.IGNORECASE
        )
        if m:
            data["pAHI"] = float(m.group(1))

    m = re.search(
        r"ODI\s*4%\s*:?\s*(?:\d+\s+)?(\d+\.\d+)\s+(\d+\.\d+)\s+(\d+\.\d+)",
        clean_text, re.IGNORECASE
    )
    if m:
        data["ODI"] = float(m.group(3))

    m = re.search(
        r"Oxygen Saturation Statistics\s+Mean\s*:\s*Minimum\s*:\s*Maximum\s*:\s*"
        r"(\d{2,3})\s+(\d{2,3})\s+(\d{2,3})",
        clean_text, re.IGNORECASE
    )
    if m:
        data["mean_sat"] = float(m.group(1))
        data["min_sat"] = float(m.group(2))

    m = re.search(r"Oxygen Saturation\s*<90.*?Sleep\s*%\s+(\d+(?:\.\d+)?)", clean_text, re.IGNORECASE)
    if m:
        data["sat_below_90_pct"] = float(m.group(1))

    m = re.search(r"Snoring Statistics.*?Mean\s*:?\s*(\d+(?:\.\d+)?)\s*dB", clean_text, re.IGNORECASE)
    if m:
        data["snoring_mean_db"] = float(m.group(1))

    m_pos = re.search(r'Body Position Statistics.*?Sleep %\s*([\d\.]+)\s+([\d\.]+)\s+([\d\.]+)\s+([\d\.]+)\s+([\d\.]+)', clean_text)
    if m_pos:
        data["pos_stats"] = {
            "Supine": float(m_pos.group(1)),
            "Prone": float(m_pos.group(2)),
            "Right": float(m_pos.group(3)),
            "Left": float(m_pos.group(4)),
        }

    m_pahi = re.search(r'pAHI 3%\s*:?\s*([\d\.]+|N/A)\s+([\d\.]+|N/A)\s+([\d\.]+|N/A)\s+([\d\.]+|N/A)\s+([\d\.]+|N/A)', clean_text, re.IGNORECASE)
    if m_pahi:
        def _parse_val(v):
            return float(v) if v != 'N/A' else None
        data["resp_stats"] = {
            "pAHI 3%": {
                "all": data.get("pAHI", 0.0),
                "supine": _parse_val(m_pahi.group(1)),
                "nsupine": _parse_val(m_pahi.group(5))
            }
        }
    else:
        # Fallback if positional pAHI is missing
        data["resp_stats"] = {
            "pAHI 3%": {
                "all": data["pAHI"],
                "supine": None,
                "nsupine": None
            }
        }

    return data


def generate_watchpat_summary(data: dict, lang: str = "en") -> str:
    pahi = data.get("pAHI", 0)
    tst = data.get("tst_hours", 0)
    t90 = data.get("sat_below_90_pct", 0)
    avg_sat = data.get("mean_sat", 100)
    min_sat = data.get("min_sat", 100)

    # 1. Determine the conclusion text based on the flowchart
    if tst > 0 and tst < 4.0:
        if lang == "he":
            conclusion = "לסיכום: משך השינה היה קצר, דבר העלול להפחית את מהימנות התוצאות.\nהערה: נדרשת בדיקת רופא (Send to review by a physician)."
        else:
            conclusion = "To summarize: the sleep duration was short, which may reduce the reliability of the results.\nNote: Send to review by a physician."
    elif pahi < 5:
        if t90 <= 1.0 and min_sat >= 90:
            if lang == "he":
                conclusion = "לסיכום: בדיקת השינה תקינה."
            else:
                conclusion = "To summarize: the sleep study is normal."
        elif t90 > 5.0 or avg_sat < 94 or min_sat < 85:
            if lang == "he":
                conclusion = "לסיכום: בדיקת השינה מצביעה על הפרעת נשימה בשינה בדרגה קלה, עם דה-סטורציה לילית שאינה פרופורציונלית.\nהמלצות: המשך הערכה/טיפול במסגרת מרפאת שינה."
            else:
                conclusion = "To summarize: the sleep study indicates mild sleep-disordered breathing, with disproportionate nocturnal desaturation.\nRecommendations: continue evaluation/treatment within a sleep clinic."
        else:
            if lang == "he":
                conclusion = "לסיכום: בדיקת השינה אינה מצביעה על הפרעת נשימה משמעותית בשינה לפי קריטריון AHI. עם זאת, נצפתה ירידה בריווי החמצן בדם בלילה.\nהמלצות: המשך הערכה/טיפול במסגרת מרפאת שינה."
            else:
                conclusion = "To summarize: the sleep study does not indicate significant sleep-disordered breathing according to AHI criteria. However, a drop in blood oxygen saturation was observed at night.\nRecommendations: continue evaluation/treatment within a sleep clinic."
    else:
        if pahi < 15:
            if lang == "he":
                conclusion = "לסיכום: בדיקת השינה מצביעה על הפרעת נשימה בשינה בדרגה קלה.\nהמלצות: המשך הערכה/טיפול במסגרת מרפאת שינה."
            else:
                conclusion = "To summarize: the sleep study indicates mild sleep-disordered breathing.\nRecommendations: continue evaluation/treatment within a sleep clinic."
        elif pahi <= 30:
            if lang == "he":
                conclusion = "לסיכום: בדיקת השינה מצביעה על הפרעת נשימה בשינה בדרגה בינונית.\nהמלצות: המשך הערכה/טיפול במסגרת מרפאת שינה."
            else:
                conclusion = "To summarize: the sleep study indicates moderate sleep-disordered breathing.\nRecommendations: continue evaluation/treatment within a sleep clinic."
        else:
            if lang == "he":
                conclusion = "לסיכום: בדיקת השינה מצביעה על הפרעת נשימה בשינה בדרגה חמורה.\nהמלצות: המשך הערכה/טיפול במסגרת מרפאת שינה."
            else:
                conclusion = "To summarize: the sleep study indicates severe sleep-disordered breathing.\nRecommendations: continue evaluation/treatment within a sleep clinic."

    # 2. Format snoring
    if data.get("snoring_mean_db", 0) > 45:
        if lang == "he":
            snoring_text = f"זוהו נחירות בעוצמה של {data['snoring_mean_db']:.0f} דציבל (בינונית/חמורה)."
        else:
            snoring_text = f"Snoring was detected at an intensity of {data['snoring_mean_db']:.0f} dB (Moderate/Severe)."
    elif data.get("snoring_mean_db", 0) > 0:
        if lang == "he":
            snoring_text = f"זוהו נחירות בעוצמה של {data['snoring_mean_db']:.0f} דציבל (קלה)."
        else:
            snoring_text = f"Snoring was detected at an intensity of {data['snoring_mean_db']:.0f} dB (Mild)."
    else:
        if lang == "he":
            snoring_text = "לא זוהו נחירות."
        else:
            snoring_text = "Snoring was not detected."

    # 3. Build the final string
    if lang == "he":
        summary = (
            f"תוצאות בדיקת שינה WatchPAT\n"
            f"בדיקת שינה ביתית בוצעה באמצעות מכשיר WatchPAT\n"
            f"תאריך בדיקת השינה: {data.get('date', 'N/A')}.\n"
            f"זמן ניטור שינה כולל: {tst:.1f} שעות.\n"
            f"יעילות שינה: {data.get('efficiency', 0):.1f}%. "
            f"מספר יקיצות: {data.get('wakes_per_hour', 0):.1f} לשעה.\n\n"
            f"{snoring_text}\n\n"
            f"אירועי נשימה (הפסקות נשימה וירידות בנשימה) נצפו בתדירות של {pahi:.1f} לשעה.\n"
            f"אירועים המלווים בירידה בריווי החמצן נצפו בתדירות של {data.get('ODI', 0):.1f} לשעה.\n"
            f"ריווי חמצן ממוצע: {avg_sat:.0f}%. "
            f"ריווי חמצן מינימלי: {min_sat:.0f}%.\n"
            f"ריווי חמצן מתחת ל-90% נצפה במשך {t90:.1f}% מזמן הניטור.\n\n"
            f"קשר לתנוחת שינה: (נדרש פענוח ידני של טבלת תנוחות הגוף).\n"
            f"קשר לשנת REM: (נדרש פענוח ידני של מדדי נשימה).\n\n"
            f"{conclusion}\n"
        )
    else:
        summary = (
            f"WatchPAT sleep study results\n"
            f"A home sleep study was performed using WatchPAT\n"
            f"Date of the sleep study: {data.get('date', 'N/A')}.\n"
            f"Total sleep monitoring duration: {tst:.1f} hours.\n"
            f"Sleep efficiency: {data.get('efficiency', 0):.1f}%. "
            f"Number of awakenings: {data.get('wakes_per_hour', 0):.1f} per hour.\n\n"
            f"{snoring_text}\n\n"
            f"Respiratory events (apneas and hypopneas) were observed at a frequency of {pahi:.1f} per hour.\n"
            f"Events associated with oxygen desaturation were observed at a frequency of {data.get('ODI', 0):.1f} per hour.\n"
            f"Average saturation: {avg_sat:.0f}%. "
            f"Minimum saturation: {min_sat:.0f}%.\n"
            f"Saturation below 90% for {t90:.1f}% of the monitoring time.\n\n"
            f"Association with sleep position: (Requires manual review of Body Position table).\n"
            f"Association with REM sleep: (Requires manual review of Respiratory Indices).\n\n"
            f"{conclusion}\n"
        )
    return summary



DOCX_TRANSLATIONS = {
    'Sleep Study Report - Detailed Analysis': 'דוח מחקר שינה - ניתוח מפורט',
    'Patient Information': 'מידע על המטופל',
    'Subject ID:': 'תעודת זהות:',
    'Age:': 'גיל:',
    'years': 'שנים',
    'Sex:': 'מין:',
    'Male': 'זכר',
    'Female': 'נקבה',
    'Recording Duration:': 'משך ההקלטה:',
    'Parameter': 'פרמטר',
    'Result': 'תוצאה',
    'Value': 'ערך',
    'Normal Range': 'טווח תקין',
    'Sleep Architecture': 'ארכיטקטורת שינה',
    'Total Sleep Time (min)': 'זמן שינה כולל (דקות)',
    'Sleep Latency (min)': 'חביון שינה (דקות)',
    'Sleep Efficiency': 'יעילות שינה',
    'WASO (min)': 'זמן ערות לאחר תחילת שינה (דקות)',
    'Sleep Stages Distribution': 'פיזור שלבי השינה',
    'Sleep Stage': 'שלב שינה',
    'Percentage': 'אחוזים',
    'Wake': 'ערות',
    'N1 (Light Sleep)': 'N1 (שינה קלה)',
    'N2 (Light Sleep)': 'N2 (שינה קלה)',
    'N3 (Deep/SWS)': 'N3 (שינה עמוקה)',
    'REM Sleep': 'שנת חלום (REM)',
    'Hypnogram': 'היפנוגרמה',
    'Visual representation of sleep stages throughout the recording:': 'ייצוג חזותי של שלבי השינה לאורך ההקלטה:',
    'Respiratory & Oxygenation Analysis': 'ניתוח נשימה וחמצון',
    'Oxygen Saturation (SpO2)': 'ריווי חמצן בדם (SpO2)',
    'Average SpO2': 'ממוצע SpO2',
    'Minimum SpO2': 'מינימום SpO2',
    'Time Below 90% (min)': 'זמן מתחת ל-90% (דקות)',
    'Limb Movements Analysis (AASM Criteria)': 'ניתוח תנועות גפיים (קריטריוני AASM)',
    'Total Limb Movements': 'סה"כ תנועות גפיים',
    'Limb Movement Index (LMSI)': 'אינדקס תנועות גפיים (LMSI)',
    'Periodic Limb Movements (PLM)': 'תנועות גפיים מחזוריות (PLM)',
    'PLM Index (PLMI)': 'אינדקס PLM (PLMI)',
    'WatchPAT Graphs': 'גרפים של WatchPAT',
    'WatchPAT Clinical Summary': 'סיכום קליני של WatchPAT',
    'pAHI (events/hr)': 'pAHI (אירועים/שעה)',
    'ODI 4% (events/hr)': 'ODI 4% (אירועים/שעה)',
    'Mean SpO2': 'ממוצע SpO2',
    'Time below SpO2 90%': 'זמן מתחת ל-SpO2 90%',
    'Snoring mean intensity': 'עוצמת נחירות ממוצעת',
    'Clinical Notes': 'הערות קליניות',
    'Sleep Stage Definitions:': 'הגדרות שלבי השינה:',
    '• N1: Transition from wakefulness to sleep (light sleep)': '• N1: מעבר מעירות לשינה (שינה קלה)',
    '• N2: Consolidated light sleep with sleep spindles and K-complexes': '• N2: שינה קלה מבוססת עם כישורי שינה ומכלולי K',
    '• N3 (SWS): Deep slow-wave sleep, restorative sleep stage': '• N3 (SWS): שנת גלים איטיים עמוקה, שלב שינה משקם',
    '• REM: Rapid Eye Movement sleep, associated with dreaming': '• REM: שנת תנועות עיניים מהירות, קשורה לחלימה',
    'Note: ': 'הערה: ',
    'Results marked with ▲/▼ deviate more than 2 Standard Deviations from age- and sex-matched normative data.': 'תוצאות המסומנות עם ▲/▼ חורגות ביותר משתי סטיות תקן מנתונים נורמטיביים תואמי גיל ומין.',
    'Page ': 'עמוד ',
}


def _create_docx_detailed(
    *,
    subject_id: str,
    out_dir: Path,
    hypno_1hz_int: np.ndarray,
    subject_age: int,
    sex: str,
    hypno_png_path: Path,
    pie_png_path: Path,
    spo2_edf_path: str | Path | None = None,
    spo2_channel: str = "SaO2 (SPO2)",
    pleth_channel: str = "Pulse (Pleth)",
    plm_metrics: dict = None,
    watchpat_pdf_path: str | Path | None = None,
    motor_events_png_path: Path | None = None,
    lang: str = "en",
) -> Path:


    subject_id = str(subject_id).split('_')[0]
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    
    def _t(text):
        if lang == "he":
            return DOCX_TRANSLATIONS.get(text, text)
        return text

    normal_ranges = get_result_ranges(subject_age, sex)
    hypno_stats = hypno_report_stats(hypno_1hz_int)

    avg_spo2 = None
    minimal_spo2 = None
    time_under_90 = None
    if spo2_edf_path:
        try:
            avg_spo2, minimal_spo2, time_under_90, _ = analyse_spo2(
                spo2_edf_path, spo2_channel=spo2_channel, pleth_channel=pleth_channel
            )
        except Exception:
            avg_spo2, minimal_spo2, time_under_90 = None, None, None

    watchpat_data = None
    watchpat_summary_text = None
    if watchpat_pdf_path and Path(watchpat_pdf_path).exists():
        try:
            watchpat_data = parse_watchpat_pdf(watchpat_pdf_path)
            watchpat_summary_text = generate_watchpat_summary(watchpat_data, lang=lang)
        except Exception:
            pass

    time_examination = _seconds_to_hms_str(hypno_stats["TIB"])
    tst_min  = hypno_stats["SPT"]  / 60.0
    lat_min  = hypno_stats["LAT"]  / 60.0
    waso_min = hypno_stats["WASO"] / 60.0
    time_sleep_str = f"{tst_min:.1f}"
    lat_str        = f"{lat_min:.1f}"
    waso_str       = f"{waso_min:.1f}"
    b90_str = f"{time_under_90 / 60.0:.1f}" if time_under_90 is not None else "N/A"

    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # RTL Helper
    def set_rtl(run=None, para=None, table=None):
        if lang != "he": return
        if run:
            run.font.rtl = True
        if para:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if table:
            tblPr = table._element.xpath('w:tblPr')
            if tblPr:
                bidi = OxmlElement('w:bidiVisual')
                tblPr[0].append(bidi)

    def add_p(text='', bold=False, level=None):
        if level is not None:
            p = document.add_heading(_t(text), level=level)
        else:
            p = document.add_paragraph(_t(text))
        set_rtl(para=p)
        if p.runs:
            for r in p.runs:
                if bold: r.bold = True
                set_rtl(run=r)
        return p

    def add_r(para, text, bold=False):
        r = para.add_run(_t(text))
        if bold: r.bold = True
        set_rtl(run=r)
        return r

    def _resource_path(relative_path):
        import sys, os
        if hasattr(sys, '_MEIPASS'):
            return Path(sys._MEIPASS) / relative_path
        return Path(__file__).parent / relative_path

    # --- HEADER ---
    section = document.sections[0]
    header_para = section.header.paragraphs[0]
    header_para.text = ""
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_up_path = _resource_path("logo_up.png")
    if logo_up_path.exists():
        run = header_para.add_run()
        run.add_picture(str(logo_up_path), width=Inches(3.0))

    # --- FOOTER ---
    footer_para = section.footer.paragraphs[0]
    footer_para.text = ""
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_down_path = _resource_path("logo_down.png")
    if logo_down_path.exists():
        run = footer_para.add_run()
        run.add_picture(str(logo_down_path), width=Inches(6.0))
    else:
        footer_run = footer_para.add_run(_t("Page "))
        footer_run.font.name = 'Calibri'
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_run.font.size = Pt(9)
        fldChar1 = OxmlElement('w:fldChar');  fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar');  fldChar2.set(qn('w:fldCharType'), 'end')
        footer_run._r.append(fldChar1); footer_run._r.append(instrText); footer_run._r.append(fldChar2)

    # --- TITLE ---
    title = add_p('Sleep Study Report - Detailed Analysis', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = 'Calibri'
    title.runs[0].font.color.rgb = RGBColor(29, 53, 87)

    # --- PATIENT INFO ---
    add_p('Patient Information', level=1)
    info_table = document.add_table(rows=3, cols=2)
    set_rtl(table=info_table)
    
    sex_str = 'Male' if str(sex).upper().startswith('M') else 'Female'
    
    info_table.rows[0].cells[0].text = _t('Subject ID:')
    info_table.rows[0].cells[1].text = subject_id
    info_table.rows[1].cells[0].text = _t('Age:')
    info_table.rows[1].cells[1].text = f"{subject_age} {_t('years')}"
    info_table.rows[2].cells[0].text = _t('Sex:')
    info_table.rows[2].cells[1].text = _t(sex_str)
    
    for row in info_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_rtl(para=p)
                for r in p.runs: set_rtl(run=r)
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    add_p()
    dp = add_p()
    add_r(dp, 'Recording Duration: ', bold=True)
    add_r(dp, time_examination)

    def style_table_header(table):
        for cell in table.rows[0].cells:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1D3557')
            cell._tc.get_or_add_tcPr().append(shd)
            for para in cell.paragraphs:
                set_rtl(para=para)
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    set_rtl(run=run)

    def add_row_no_ref(row_idx: int, param: str, result: str, normal: str, table):
        cells = table.rows[row_idx].cells
        cells[0].text = _t(param)
        indicator = ""
        if _check_if_abnormal(result, normal):
            vm = re.search(r"[-+]?\d*\.\d+|\d+", str(result))
            nm = re.search(r"([-+]?\d*\.\d+|\d+)\s*±", str(normal))
            if vm and nm:
                indicator = " ▲" if float(vm.group()) > float(nm.group(1)) else " ▼"
        
        cells[1].text = ""
        p = cells[1].paragraphs[0]
        run = p.add_run(f"{result}{indicator}")
        if indicator:
            run.font.color.rgb = RGBColor(230, 57, 70)
            run.font.bold = True
            
        for cell in cells:
            for para in cell.paragraphs:
                set_rtl(para=para)
                for r in para.runs: set_rtl(run=r)

    # 1: SLEEP ARCHITECTURE
    add_p('Sleep Architecture', level=1)
    t = document.add_table(rows=5, cols=2)
    set_rtl(table=t)
    t.style = 'Light Grid Accent 1'
    t.rows[0].cells[0].text = _t('Parameter'); t.rows[0].cells[1].text = _t('Result')
    style_table_header(t)
    add_row_no_ref(1, 'Total Sleep Time (min)', time_sleep_str, normal_ranges['TST'], t)
    add_row_no_ref(2, 'Sleep Latency (min)',    lat_str,        normal_ranges['LAT'], t)
    add_row_no_ref(3, 'Sleep Efficiency',       f"{hypno_stats['SE']:.1f}%", normal_ranges['SE'],  t)
    add_row_no_ref(4, 'WASO (min)',             waso_str,       normal_ranges['WASO'], t)

    add_p()
    add_p('Sleep Stages Distribution', level=1)
    st = document.add_table(rows=6, cols=2)
    set_rtl(table=st)
    st.style = 'Light Grid Accent 1'
    st.rows[0].cells[0].text = _t('Sleep Stage'); st.rows[0].cells[1].text = _t('Percentage')
    style_table_header(st)
    add_row_no_ref(1, 'Wake',             f"{hypno_stats['Wake']:.1f}%", 'N/A',                st)
    add_row_no_ref(2, 'N1 (Light Sleep)', f"{hypno_stats['N1']:.1f}%",   normal_ranges['N1'],  st)
    add_row_no_ref(3, 'N2 (Light Sleep)', f"{hypno_stats['N2']:.1f}%",   normal_ranges['N2'],  st)
    add_row_no_ref(4, 'N3 (Deep/SWS)',   f"{hypno_stats['N3']:.1f}%",   normal_ranges['N3'],  st)
    add_row_no_ref(5, 'REM Sleep',        f"{hypno_stats['REM']:.1f}%",  normal_ranges['REM'], st)

    add_p()
    add_p('Hypnogram', level=1)
    add_p('Visual representation of sleep stages throughout the recording:')
    if Path(hypno_png_path).exists():
        document.add_picture(str(hypno_png_path), width=Inches(6.5))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_p()
    if Path(pie_png_path).exists():
        document.add_picture(str(pie_png_path), width=Inches(5.0))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2: RESPIRATORY
    add_p()
    add_p('Respiratory & Oxygenation Analysis', level=1)
    if avg_spo2 is not None:
        add_p('Oxygen Saturation (SpO2)', level=2)
        spo2t = document.add_table(rows=4, cols=2)
        set_rtl(table=spo2t)
        spo2t.style = 'Light Grid Accent 1'
        spo2t.rows[0].cells[0].text = _t('Parameter'); spo2t.rows[0].cells[1].text = _t('Result')
        style_table_header(spo2t)
        add_row_no_ref(1, 'Average SpO2',        f"{avg_spo2:.2f}%",        normal_ranges['AvS'],  spo2t)
        add_row_no_ref(2, 'Minimum SpO2',         f"{minimal_spo2[1]:.2f}%", normal_ranges['MinS'], spo2t)
        add_row_no_ref(3, 'Time Below 90% (min)', b90_str,                   normal_ranges['B90'],  spo2t)
        add_p()

    if plm_metrics is not None and plm_metrics["Total_LMS"] > 0:
        add_p('Limb Movements Analysis (AASM Criteria)', level=2)
        plmt = document.add_table(rows=5, cols=2)
        set_rtl(table=plmt)
        plmt.style = 'Light Grid Accent 1'
        plmt.rows[0].cells[0].text = _t('Parameter'); plmt.rows[0].cells[1].text = _t('Result')
        style_table_header(plmt)
        add_row_no_ref(1, 'Total Limb Movements',          str(plm_metrics["Total_LMS"]),      'N/A',    plmt)
        add_row_no_ref(2, 'Limb Movement Index (LMSI)',    f"{plm_metrics['LMSI']:.1f} / hr",  'N/A',    plmt)
        add_row_no_ref(3, 'Periodic Limb Movements (PLM)', str(plm_metrics["Total_PLM"]),      'N/A',    plmt)
        add_row_no_ref(4, 'PLM Index (PLMI)',              f"{plm_metrics['PLMI']:.1f} / hr",  '< 15.0', plmt)
        add_p()

    if watchpat_pdf_path and Path(watchpat_pdf_path).exists():
        add_p('WatchPAT Graphs', level=2)
        try:
            for img_path in extract_watchpat_visuals(Path(watchpat_pdf_path), out_dir):
                if img_path.exists():
                    document.add_picture(str(img_path), width=Inches(6.2))
                    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_p()
        except Exception:
            pass

    if motor_events_png_path and motor_events_png_path.exists():
        try:
            document.add_picture(str(motor_events_png_path), width=Inches(6.2))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_p()
        except Exception:
            pass

    if watchpat_pdf_path and Path(watchpat_pdf_path).exists():
        add_p()
        add_p('WatchPAT Clinical Summary', level=1)
        if watchpat_data is not None:
            wpt = document.add_table(rows=7, cols=2)
            set_rtl(table=wpt)
            wpt.style = 'Light Grid Accent 1'
            wpt.rows[0].cells[0].text = _t('Parameter'); wpt.rows[0].cells[1].text = _t('Value')
            style_table_header(wpt)
            for i, (param, val) in enumerate([
                ('pAHI (events/hr)',       f"{watchpat_data['pAHI']:.1f}"),
                ('ODI 4% (events/hr)',      f"{watchpat_data['ODI']:.1f}"),
                ('Mean SpO2',               f"{watchpat_data['mean_sat']:.0f}%"),
                ('Minimum SpO2',            f"{watchpat_data['min_sat']:.0f}%"),
                ('Time below SpO2 90%',     f"{watchpat_data['sat_below_90_pct']:.1f}%"),
                ('Snoring mean intensity',  f"{watchpat_data['snoring_mean_db']:.0f} dB"),
            ], start=1):
                wpt.rows[i].cells[0].text = _t(param)
                wpt.rows[i].cells[1].text = val
                for cell in wpt.rows[i].cells:
                    for para in cell.paragraphs:
                        set_rtl(para=para)
                        for r in para.runs: set_rtl(run=r)
            add_p()

        if watchpat_summary_text:
            for i, line in enumerate(watchpat_summary_text.split('\n')):
                line = line.strip()
                if not line:
                    continue
                # Assuming this text is generated in English, we might not have direct translations for it.
                # Just add it as is, but apply RTL if needed.
                p = add_p(line)
                if i < 2 and p.runs:
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(11)

    add_p()
    add_p('Clinical Notes', level=1)
    dp = add_p()
    add_r(dp, 'Sleep Stage Definitions:', bold=True)
    add_p('• N1: Transition from wakefulness to sleep (light sleep)')
    add_p('• N2: Consolidated light sleep with sleep spindles and K-complexes')
    add_p('• N3 (SWS): Deep slow-wave sleep, restorative sleep stage')
    add_p('• REM: Rapid Eye Movement sleep, associated with dreaming')
    add_p()
    dp = add_p()
    add_r(dp, 'Note: ', bold=True)
    add_r(dp, 'Results marked with ▲/▼ deviate more than 2 Standard Deviations from age- and sex-matched normative data.')

    suffix = f"_{lang}" if lang != "en" else ""
    out_path = out_dir / f"sleep_report_{subject_id}_detailed{suffix}.docx"
    document.save(str(out_path))
    return out_path
def create_docx_sleep_report(
    *,
    subject_id: str,
    outdir: Path,
    hypno_1hz_int: np.ndarray,
    subject_age: int,
    sex: str,
    hypno_png_path: Path,
    pie_png_path: Path,
    spo2_edf_path: str | Path | None = None,
    spo2_channel: str = "SaO2 SPO2",
    pleth_channel: str = "Pulse Pleth",
    report_type: str = "simple",
    plm_metrics: dict = None,
    watchpat_pdf_path: str | Path | None = None, # Добавлен аргумент
    motor_events_png_path: Path | None = None,
    lang: str = "en",
    subject_name: str | None = None,
) -> Path:
    from docx import Document
    from docx.shared import Inches
    if subject_name:
        subject_id = subject_name
    else:
        subject_id = str(subject_id).split('_')[0]
    outdir = Path(outdir)
    outdir.mkdir(parents=True, exist_ok=True)



    if report_type == "detailed":
        out_en = _create_docx_detailed(
            subject_id=subject_id,
            out_dir=outdir,
            hypno_1hz_int=hypno_1hz_int,
            subject_age=subject_age,
            sex=sex,
            hypno_png_path=hypno_png_path,
            pie_png_path=pie_png_path,
            spo2_edf_path=spo2_edf_path,
            spo2_channel=spo2_channel,
            pleth_channel=pleth_channel,
            plm_metrics=plm_metrics,
            watchpat_pdf_path=watchpat_pdf_path,
            motor_events_png_path=motor_events_png_path,
            lang="en"
        )
        out_he = None
        try:
            out_he = _create_docx_detailed(
                subject_id=subject_id,
                out_dir=outdir,
                hypno_1hz_int=hypno_1hz_int,
                subject_age=subject_age,
                sex=sex,
                hypno_png_path=hypno_png_path,
                pie_png_path=pie_png_path,
                spo2_edf_path=spo2_edf_path,
                spo2_channel=spo2_channel,
                pleth_channel=pleth_channel,
                plm_metrics=plm_metrics,
                watchpat_pdf_path=watchpat_pdf_path,
                motor_events_png_path=motor_events_png_path,
                lang="he"
            )
        except Exception as e:
            print(f"[Warning] Failed to generate HE detailed docx: {e}")
        return out_en, out_he


    normal_ranges = get_result_ranges(subject_age, sex)
    hypno_stats = hypno_report_stats(hypno_1hz_int)

    avg_spo2 = None
    minimal_spo2 = None
    time_under_90 = None
    if spo2_edf_path:
        try:
            avg_spo2, minimal_spo2, time_under_90, _ = analyse_spo2(
                spo2_edf_path,
                spo2_channel=spo2_channel,
                pleth_channel=pleth_channel,
            )
        except Exception:
            avg_spo2, minimal_spo2, time_under_90 = None, None, None

    document = Document()
    document.add_heading('דו"ח שינה', 0)

    time_examination = _seconds_to_hms_str(hypno_stats["TIB"])
    
    tst_min = hypno_stats["SPT"] / 60.0
    lat_min = hypno_stats["LAT"] / 60.0
    waso_min = hypno_stats["WASO"] / 60.0
    
    time_sleep_str = f"{tst_min:.1f}"
    lat_str = f"{lat_min:.1f}"
    waso_str = f"{waso_min:.1f}"
    
    if time_under_90 is not None:
        b90_min = time_under_90 / 60.0
        b90_str = f"{b90_min:.1f}"
    else:
        b90_str = "N/A"

    document.add_paragraph(" תאריך הבדיקה:")
    p = document.add_paragraph("שם הנבדק/ת: ")
    p.add_run(subject_id)
    p = document.add_paragraph('משך הבדיקה: ')
    p.add_run(time_examination)
    document.add_paragraph('דוח זה הופק אוטומטית מתוך היפנוגרמה, ספקטרוגרמה ומדדי שינה מחוшבים.')

    table_characteristics = document.add_table(rows=1, cols=3)
    table_characteristics.style = "Table Grid"
    hdr_cells = table_characteristics.rows[0].cells
    hdr_cells[0].text = 'תוצאה אופיינית בקבוצת הגיל'
    hdr_cells[1].text = 'תוצאה'
    hdr_cells[2].text = 'מאפיין שינה '

    def add_char_row(ref_text: str, result_text: str, label: str):
        row_cells = table_characteristics.add_row().cells
        row_cells[0].text = ref_text
        row_cells[1].text = result_text
        row_cells[2].text = label

    add_char_row(normal_ranges["TST"], time_sleep_str, "זמן השינה הכולל (דקות)")
    add_char_row(normal_ranges["LAT"], lat_str, "זמן עד להרדמות (דקות)")
    add_char_row(normal_ranges["SE"], f"{hypno_stats['SE']:.2f}%", "יעילות השינה")
    add_char_row(normal_ranges["WASO"], waso_str, "זמן בעירות לאחר הירדמות ראשונית (דקות)")
    add_char_row(normal_ranges["AvS"], f"{avg_spo2:.2f}%" if avg_spo2 is not None else "N/A", "סטורציה ממוצעת")
    add_char_row(normal_ranges["MinS"], f"{minimal_spo2[1]:.2f}%" if minimal_spo2 is not None else "N/A", "סטורציה מינימלית")
    add_char_row(normal_ranges["B90"], b90_str, "זמן בסטורציה נמוכה מ90% (דקות)")
    
    if plm_metrics is not None and plm_metrics["Total_LMS"] > 0:
        add_char_row("< 15.0", f"{plm_metrics['PLMI']:.1f}", "מדד תנועות רגליים מחזוריות (PLMI)")
        add_char_row("N/A", str(plm_metrics["Total_PLM"]), "סה״כ תנועות רגליים מחזוריות בשינה")

    # Вставка визуалов в простую версию отчета (один за другим)
    if watchpat_pdf_path and Path(watchpat_pdf_path).exists():
        watchpat_imgs = extract_watchpat_visuals(Path(watchpat_pdf_path), outdir)
        if watchpat_imgs:
            document.add_heading('נתוני WatchPAT (סטורציה ודופק)', level=1)
            for img_path in watchpat_imgs:
                if img_path.exists():
                    document.add_picture(str(img_path), width=Inches(5.5))
                    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading('דיאגרמת שלבי השינה - היפנוגרמה: ', level=1)
    if Path(hypno_png_path).exists():
        document.add_picture(str(hypno_png_path), width=Inches(5.25))

    document.add_heading('התפלגות שלבי השינה:', level=1)
    table_sleep = document.add_table(rows=1, cols=3)
    table_sleep.style = "Table Grid"
    hdr_cells = table_sleep.rows[0].cells
    hdr_cells[0].text = 'תוצאה ממוצעת אצל קבוצת הגיל'
    hdr_cells[1].text = 'תוצאות'
    hdr_cells[2].text = 'חלוקת השינה לשלבי השינה באחוזים'

    def add_stage_row(ref_text: str, result_text: str, label: str):
        row_cells = table_sleep.add_row().cells
        row_cells[0].text = ref_text
        row_cells[1].text = result_text
        row_cells[2].text = label

    add_stage_row(normal_ranges["N1"], f"{hypno_stats['N1']:.2f}%", "שלб השינה הראשון N1 שינה קלה")
    add_stage_row(normal_ranges["N2"], f"{hypno_stats['N2']:.2f}%", "שלב השינה השני N2")
    add_stage_row(normal_ranges["N3"], f"{hypno_stats['N3']:.2f}%", "שלб השינה השלישי N3 שינה עמוקה")
    add_stage_row(normal_ranges["REM"], f"{hypno_stats['REM']:.2f}%", "שנת חלום (REM Sleep)")

    if Path(pie_png_path).exists():
        document.add_picture(str(pie_png_path), width=Inches(5.25))

    out_path = outdir / f"sleep_report_{subject_id}.docx"
    document.save(str(out_path))
    return out_path

def _run_rbdtector_headless_case(
    *,
    outdir: Path,
    raw_for_edf: mne.io.BaseRaw,
    hypno_1hz_int: np.ndarray,
    emg_label_in_edf: str,
    verbose: bool = True,
) -> dict[str, Any]:
    case_dir = Path(outdir) / "rbdtector_case"
    case_dir.mkdir(parents=True, exist_ok=True)

    start_dt = _choose_start_dt(raw_for_edf)
    edf_path = case_dir / "export.edf"

    rr = raw_for_edf.copy()
    _ensure_loaded(rr)
    rr_fixed = rr.copy()

    rr_fixed.set_meas_date(None)
    
    try:
        with rr_fixed.info._unlock():
            rr_fixed.info['device_info'] = None
            rr_fixed.info['subject_info'] = None
            rr_fixed.info['experimenter'] = None
            rr_fixed.info['proj_name'] = None
    except Exception:
        pass

    mne.export.export_raw(
        str(edf_path),
        rr_fixed,
        fmt="edf",
        overwrite=True,
        physical_range="channelwise",
    )
    assert edf_path.exists()

    sleep_profile_path = case_dir / "Sleep profile.txt"
    flow_events_path = case_dir / "Flow Events.txt"
    arousals_path = case_dir / "Classification Arousals.txt"

    _write_sleep_profile_txt(
        sleep_profile_path,
        start_dt=start_dt,
        hypno_1hz_int=hypno_1hz_int,
    )
    _write_minimal_event_file(
        flow_events_path,
        start_dt=start_dt,
        label="FLOW; NONE",
    )
    _write_minimal_event_file(
        arousals_path,
        start_dt=start_dt,
        label="AROUSAL; NONE",
    )

    if emg_label_in_edf not in rr_fixed.ch_names:
        raise ValueError(
            f"EMG channel '{emg_label_in_edf}' not found in EDF export channels: {rr_fixed.ch_names}"
        )

    settings.SIGNALS_TO_EVALUATE = [emg_label_in_edf]
    settings.CHIN = 0
    settings.LEGS = False
    settings.ARMS = False
    settings.FLOW = False
    settings.HUMAN_ARTIFACTS = False
    settings.SNORE = False
    settings.HUMAN_BASELINE = False
    # ЖЕСТКАЯ ФИЛЬТРАЦИЯ КОРОТКИХ СПАЙКОВ (ЭКГ)
    settings.MIN_SUSTAINED = 0.25 

    from data_structures.raw_data_channel import RawDataChannel
    from app_logic.PSG import PSG
    import pandas as pd
    import numpy as np

    old_get_sample_rate = RawDataChannel.get_sample_rate
    old_create_sleep_profile_column = PSG.create_sleep_profile_column

    def _patched_get_sample_rate(self):
        hdr = getattr(self, "_signal_header", None) or {}
        for key in ("sample_rate", "sample_frequency", "sampling_rate", "sfreq"):
            val = hdr.get(key, None)
            if val is not None:
                try:
                    return float(val)
                except Exception:
                    return val
        if verbose:
            print(f"DEBUG signal header keys: {list(hdr.keys())}")
        raise KeyError(f"sample rate not found in signal header keys: {list(hdr.keys())}")

    @staticmethod
    def _patched_create_sleep_profile_column(idx: pd.DatetimeIndex, annotation_data):
        from util.definitions import SLEEP_CLASSIFIERS

        def _map_stage(v):
            try:
                iv = int(v)
            except Exception:
                return SLEEP_CLASSIFIERS["artifact"]

            if iv == 4:
                return SLEEP_CLASSIFIERS["REM"]
            if iv == 3:
                return "N3"
            if iv == 2:
                return "N2"
            if iv == 1:
                return "N1"
            if iv == 0:
                return "wake"
            return SLEEP_CLASSIFIERS["artifact"]

        df = pd.DataFrame(index=idx)

        try:
            sleep_profile = annotation_data.sleep_profile[1].copy()
            sleep_profile.sort_index(inplace=True)
            sleep_profile = sleep_profile.loc[idx[0]:idx[-1]]
        except Exception:
            sleep_profile = pd.Series(dtype=object)

        if sleep_profile is not None and len(sleep_profile) > 0:
            try:
                sleep_profile.iloc[-1] = SLEEP_CLASSIFIERS["artifact"]
                resampled_sleep_profile = sleep_profile.resample(
                    str(1000 / settings.RATE) + "ms"
                ).ffill()

                df = pd.concat([df, resampled_sleep_profile], axis=1, join="inner")

                if "sleep_phase" in df.columns and len(df) > 0:
                    df["is_REM"] = (
                        df["sleep_phase"].astype(str).str.lower()
                        == SLEEP_CLASSIFIERS["REM"].lower()
                    )
                    if settings.SNORE:
                        df["is_SNORE"] = (
                            df["sleep_phase"].astype(str).str.lower()
                            == SLEEP_CLASSIFIERS["SNORE"].lower()
                        )
                        df["is_REM"] = np.logical_or(df["is_REM"], df["is_SNORE"])

                    if verbose:
                        rem_runs = (
                            df["is_REM"].astype(int)
                            .groupby(df["is_REM"].ne(df["is_REM"].shift()).cumsum())
                            .sum()
                        )
                        rem_runs = rem_runs[rem_runs > 0]
                        longest_sec = (
                            float(rem_runs.max()) / float(settings.RATE)
                            if len(rem_runs)
                            else 0.0
                        )
                        n_ge_150 = (
                            int((rem_runs >= 150 * settings.RATE).sum())
                            if len(rem_runs)
                            else 0
                        )
                        print(
                            "Using annotation_data.sleep_profile. "
                            f"Longest REM run: {longest_sec:.1f} sec; "
                            f"REM runs >=150 sec: {n_ge_150}"
                        )

                    return df["sleep_phase"], df["is_REM"]
            except Exception:
                pass

        hyp_1hz = np.asarray(hypno_1hz_int, dtype=int).ravel()
        if hyp_1hz.size == 0:
            raise ValueError("Fallback sleep profile cannot be created: empty hypnogram.")

        rate = int(settings.RATE)
        hyp_up = np.repeat(hyp_1hz, rate)

        n = min(len(idx), len(hyp_up))
        if n == 0:
            raise ValueError("Fallback sleep profile cannot be created: empty idx or hypnogram.")

        fallback_idx = idx[:n]
        hyp_up = hyp_up[:n]

        fallback_sleep_phase = pd.Series(
            [_map_stage(v) for v in hyp_up],
            index=fallback_idx,
            name="sleep_phase",
        )
        fallback_is_rem = pd.Series(
            hyp_up == 4,
            index=fallback_idx,
            name="is_REM",
        )

        if settings.SNORE:
            fallback_is_snore = (
                fallback_sleep_phase.astype(str).str.lower()
                == SLEEP_CLASSIFIERS["SNORE"].lower()
            )
            fallback_is_rem = np.logical_or(fallback_is_rem, fallback_is_snore)

        if verbose:
            rem_runs = (
                pd.Series(fallback_is_rem, index=fallback_idx)
                .astype(int)
                .groupby(pd.Series(fallback_is_rem, index=fallback_idx).ne(
                    pd.Series(fallback_is_rem, index=fallback_idx).shift()
                ).cumsum())
                .sum()
            )
            rem_runs = rem_runs[rem_runs > 0]
            longest_sec = (
                float(rem_runs.max()) / float(settings.RATE)
                if len(rem_runs)
                else 0.0
            )
            n_ge_150 = (
                int((rem_runs >= 150 * settings.RATE).sum())
                if len(rem_runs)
                else 0
            )
            print(
                "Using fallback sleep profile built from hypno_1hz_int "
                f"and upsampled to {settings.RATE} Hz. "
                f"Longest REM run: {longest_sec:.1f} sec; "
                f"REM runs >=150 sec: {n_ge_150}"
            )

        return fallback_sleep_phase, pd.Series(
            fallback_is_rem,
            index=fallback_idx,
            name="is_REM",
        )

    RawDataChannel.get_sample_rate = _patched_get_sample_rate
    PSG.create_sleep_profile_column = _patched_create_sleep_profile_column

    try:
        df_out, df_channel_combinations = PSGController.run_rbd_detection(
            str(case_dir),
            str(case_dir),
        )
    finally:
        RawDataChannel.get_sample_rate = old_get_sample_rate
        PSG.create_sleep_profile_column = old_create_sleep_profile_column

    df_out_csv = Path(outdir) / "rbdtector_results.csv"
    df_channel_csv = Path(outdir) / "rbdtector_channel_combinations.csv"

    if hasattr(df_out, "to_csv"):
        df_out.to_csv(df_out_csv)
    if hasattr(df_channel_combinations, "to_csv"):
        df_channel_combinations.to_csv(df_channel_csv, index=False)

    calculated_df = None
    if isinstance(df_out, pd.DataFrame) and not df_out.empty:
        calculated_df = df_out.copy()
        if not isinstance(calculated_df.index, pd.DatetimeIndex):
            calculated_df.index = pd.date_range(
                start=start_dt, 
                periods=len(calculated_df), 
                freq=f"{1000/settings.RATE}ms"
            )

    return {
        "case_dir": str(case_dir),
        "edf_path": str(edf_path),
        "sleep_profile_path": str(sleep_profile_path),
        "flow_events_path": str(flow_events_path),
        "arousals_path": str(arousals_path),
        "rbdtector_results_csv": str(df_out_csv),
        "rbdtector_channel_combinations_csv": str(df_channel_csv),
        "calculated_df": calculated_df,
        "start_dt": start_dt
    }



def _hypno_1hz_to_30s_int(hypno_1hz_int: np.ndarray, epoch_sec: int = 30) -> np.ndarray:
    hypno_1hz_int = np.asarray(hypno_1hz_int, dtype=int)

    if hypno_1hz_int.ndim != 1:
        hypno_1hz_int = hypno_1hz_int.ravel()

    n_epochs = len(hypno_1hz_int) // epoch_sec
    if n_epochs == 0:
        return np.array([], dtype=int)

    valid_stage_order = [0, 1, 2, 3, 4]

    def aggregate_epoch(epoch_vals: np.ndarray) -> int:
        epoch_vals = np.asarray(epoch_vals, dtype=int)
        epoch_vals = epoch_vals[np.isin(epoch_vals, valid_stage_order)]
        if epoch_vals.size == 0:
            return -1

        counts = np.bincount(epoch_vals, minlength=max(valid_stage_order) + 1)
        max_count = counts.max()
        winners = np.flatnonzero(counts == max_count)

        for st in (4, 3, 2, 1, 0):
            if st in winners:
                return int(st)

        return int(winners[0])

    epochs = hypno_1hz_int[: n_epochs * epoch_sec].reshape(-1, epoch_sec)
    return np.array([aggregate_epoch(epoch) for epoch in epochs], dtype=int)


def _run_yasa_art_detect(
    raw_eeg: mne.io.BaseRaw,
    hypno_1hz_int: np.ndarray,
    *,
    epoch_sec: int = 30,
    include: tuple[int, ...] = (0, 1, 2, 3, 4),
    window: int = 5,
    method: str = "covar",
    threshold: float = 3.0,
    verbose: bool = True,
) -> dict[str, Any]:
    raw_eeg = raw_eeg.copy()
    _ensure_loaded(raw_eeg)

    hypno_30 = _hypno_1hz_to_30s_int(hypno_1hz_int, epoch_sec=epoch_sec)

    sf = float(raw_eeg.info["sfreq"])
    max_epochs_data = int(raw_eeg.n_times // int(round(sf * epoch_sec)))
    n_epochs = min(len(hypno_30), max_epochs_data)

    if n_epochs == 0:
        return {
            "epoch_mask_keep": np.array([], dtype=bool),
            "epoch_mask_bad": np.array([], dtype=bool),
            "art_epochs": np.array([], dtype=bool),
            "zscores": np.array([], dtype=float),
            "hypno_30": np.array([], dtype=int),
            "n_bad": 0,
            "n_total": 0,
        }

    hypno_30 = hypno_30[:n_epochs]
    n_samp = int(round(n_epochs * epoch_sec * sf))
    data = raw_eeg.get_data()[:, :n_samp]
    
    import logging
    logging.getLogger("yasa").setLevel(logging.ERROR)

    try:
        art_epochs, zscores = yasa.art_detect(
            data,
            sf=sf,
            window=window,
            method=method,
            threshold=threshold,
        )
    except TypeError:
        art_epochs, zscores = yasa.art_detect(
            data,
            sf=sf,
            window=window,
        )

    art_epochs = np.asarray(art_epochs, dtype=bool)
    zscores = np.asarray(zscores, dtype=float)

    if art_epochs.ndim > 1:
        art_epochs = np.squeeze(art_epochs)
    if zscores.ndim > 1:
        zscores = np.squeeze(zscores)

    n = min(len(hypno_30), len(art_epochs), len(zscores))
    hypno_30 = hypno_30[:n]
    art_epochs = art_epochs[:n]
    zscores = zscores[:n]

    epoch_mask_bad = art_epochs.copy()
    epoch_mask_keep = ~epoch_mask_bad

    return {
        "epoch_mask_keep": epoch_mask_keep,
        "epoch_mask_bad": epoch_mask_bad,
        "art_epochs": art_epochs,
        "zscores": zscores,
        "hypno_30": hypno_30,
        "n_bad": int(epoch_mask_bad.sum()),
        "n_total": int(len(epoch_mask_bad)),
    }



def _save_artifact_epochs_csv(
    path: Path,
    *,
    hypno_30: np.ndarray,
    bad_mask: np.ndarray,
    zscores: np.ndarray,
) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)

    hypno_30 = np.asarray(hypno_30, dtype=int)
    bad_mask = np.asarray(bad_mask, dtype=bool)
    zscores = np.asarray(zscores, dtype=float)

    n = min(len(hypno_30), len(bad_mask), len(zscores))
    stage_map = {0: "W", 1: "N1", 2: "N2", 3: "N3", 4: "REM", -1: "ART"}

    df = pd.DataFrame({
        "epoch_idx": np.arange(n, dtype=int),
        "stage_int": hypno_30[:n],
        "stage": [stage_map.get(int(v), str(int(v))) for v in hypno_30[:n]],
        "yasa_zscore": zscores[:n],
        "is_bad": bad_mask[:n].astype(int),
    })
    df.to_csv(path, index=False)
    return path

def _compute_sleep_stats(hypno_1hz_int: np.ndarray) -> dict[str, float]:
    hypno_1hz_int = np.asarray(hypno_1hz_int, dtype=int)
    stats = yasa.sleep_statistics(hypno_1hz_int, sf_hyp=1)
    keys = ["TIB", "TST", "SE", "SOL", "WASO", "%N1", "%N2", "%N3", "%NREM", "%REM"]
    return {k: float(stats[k]) for k in keys}


def _format_sleep_stats(stats: dict[str, float]) -> dict[str, str]:
    return {k: f"{v:.2f}" for k, v in stats.items()}


def _make_stage_epochs(
    raw_eeg: mne.io.BaseRaw,
    hypno_1hz_int: np.ndarray,
    epoch_sec: int = 30,
    epoch_keep_mask: np.ndarray | None = None,
) -> tuple[mne.Epochs, dict[str, int]]:
    raw_eeg = raw_eeg.copy()
    _ensure_loaded(raw_eeg)

    hypno_1hz_int = np.asarray(hypno_1hz_int, dtype=int)
    n_epochs_hyp = len(hypno_1hz_int) // epoch_sec
    if n_epochs_hyp == 0:
        raise ValueError("Hypnogram is too short.")

    hypno_30 = hypno_1hz_int[: n_epochs_hyp * epoch_sec].reshape(-1, epoch_sec)[:, 0]
    sfreq = raw_eeg.info["sfreq"]
    total_sec_data = raw_eeg.n_times / sfreq
    n_epochs_data = int(total_sec_data // epoch_sec)
    n_epochs = min(n_epochs_hyp, n_epochs_data)
    if n_epochs == 0:
        raise ValueError("No overlapping 30s epochs between raw and hypnogram.")

    hypno_30 = hypno_30[:n_epochs]
    events = mne.make_fixed_length_events(raw_eeg, id=999, duration=epoch_sec)[:n_epochs]

    hypno_to_event = {0: 1, 1: 2, 2: 3, 3: 4, 4: 5}
    keep_mask = np.isin(hypno_30, list(hypno_to_event.keys()))

    if epoch_keep_mask is not None:
        epoch_keep_mask = np.asarray(epoch_keep_mask, dtype=bool)[:n_epochs]
        keep_mask = keep_mask & epoch_keep_mask

    events = events[keep_mask]
    hypno_30_keep = hypno_30[keep_mask]
    events[:, 2] = np.array([hypno_to_event[int(v)] for v in hypno_30_keep], dtype=int)

    event_id = {"W": 1, "N1": 2, "N2": 3, "N3": 4, "REM": 5}
    epochs = mne.Epochs(
        raw_eeg,
        events=events,
        event_id=event_id,
        tmin=0,
        tmax=epoch_sec,
        baseline=None,
        preload=True,
        on_missing="ignore",
        reject_by_annotation=False,
        verbose=False,
    )
    return epochs, event_id



def _compute_stage_psd(epochs: mne.Epochs, separate_nrem_power_spectrum: bool = True, fmin: float = 0.5, fmax: float = 40.0) -> tuple[list[str], dict[str, tuple[np.ndarray, np.ndarray]]]:
    stage_order = ["W", "N2", "N3", "REM"] if separate_nrem_power_spectrum else ["W", "NREM", "REM"]
    out: dict[str, tuple[np.ndarray, np.ndarray]] = {}

    for stage in stage_order:
        if stage == "NREM":
            arrs = []
            for st in ("N2", "N3"):
                if len(epochs[st]) > 0:
                    arrs.append(epochs[st].get_data(copy=True))
            if not arrs:
                continue
            data = np.concatenate(arrs, axis=0)
        else:
            if len(epochs[stage]) == 0:
                continue
            data = epochs[stage].get_data(copy=True)

        psds, freqs = psd_array_multitaper(data, sfreq=epochs.info["sfreq"], fmin=fmin, fmax=fmax, n_jobs=1)
        psd_mean = psds.mean(axis=0).mean(axis=0)
        psd_db = 10.0 * np.log10(psd_mean)
        out[stage] = (freqs, psd_db)
    return stage_order, out


def _make_psd_stats_panel(
    *,
    raw: mne.io.BaseRaw,
    hypno_1hz_int: np.ndarray,
    subject_id: str,
    outdir: Path,
    eeg_channel_for_report: str,
    separate_nrem_power_spectrum: bool = True,
    use_yasa_art_detect: bool = True,
    yasa_art_epoch_sec: int = 30,
    yasa_art_window: int = 5,
    yasa_art_method: str = "covar",
    yasa_art_threshold: float = 3.0,
    yasa_art_include: tuple[int, ...] = (0, 1, 2, 3, 4),
    verbose: bool = True,
) -> dict[str, Any]:
    outdir.mkdir(parents=True, exist_ok=True)
    raw_one = raw.copy().pick([eeg_channel_for_report])
    _ensure_loaded(raw_one)

    hypno_1hz_int = np.asarray(hypno_1hz_int, dtype=int)
    stats = _compute_sleep_stats(hypno_1hz_int)
    stats_fmt = _format_sleep_stats(stats)

    hypno_png = outdir / f"{subject_id}_hypno.png"
    stats_csv = outdir / f"{subject_id}_stats.csv"
    panel_png = outdir / f"{subject_id}_psd_stats.png"
    art_csv = outdir / f"{subject_id}_yasa_artifacts.csv"

    _save_hypnogram_png(hypno_png, hypno_1hz_int, title=subject_id)
    _save_kv_csv(stats_csv, stats_fmt)

    art_result = None
    epoch_keep_mask = None

    if use_yasa_art_detect:
        art_result = _run_yasa_art_detect(
            raw_one,
            hypno_1hz_int,
            epoch_sec=yasa_art_epoch_sec,
            include=tuple(yasa_art_include),
            window=yasa_art_window,
            method=yasa_art_method,
            threshold=yasa_art_threshold,
            verbose=verbose,
        )
        epoch_keep_mask = art_result["epoch_mask_keep"]

        _save_artifact_epochs_csv(
            art_csv,
            hypno_30=art_result["hypno_30"],
            bad_mask=art_result["epoch_mask_bad"],
            zscores=art_result["zscores"],
        )

    epochs, _ = _make_stage_epochs(
        raw_one,
        hypno_1hz_int,
        epoch_sec=30,
        epoch_keep_mask=epoch_keep_mask,
    )

    stage_order, stage_psd = _compute_stage_psd(
        epochs,
        separate_nrem_power_spectrum=separate_nrem_power_spectrum,
        fmin=0.5,
        fmax=40.0,
    )

    fig = plt.figure(constrained_layout=True, figsize=(14, 6))
    gs = fig.add_gridspec(1, 2, width_ratios=[1.4, 1.0])

    ax1 = fig.add_subplot(gs[0, 0])
    stage_color = {"W": "b", "N2": "g", "N3": "m", "REM": "r", "NREM": "k"}
    stage_label = {"W": "wake", "N2": "N2", "N3": "N3", "REM": "REM", "NREM": "NREM"}

    for stage in stage_order:
        if stage in stage_psd:
            freqs, psd_db = stage_psd[stage]
            ax1.plot(freqs, psd_db, color=stage_color[stage], label=stage_label[stage], linewidth=2)

    ax1.set_title(f"PSD according to stages - {subject_id}")
    ax1.set_ylabel("Power Spectral Density (dB)")
    ax1.set_xlabel("Frequency (Hz)")
    if ax1.lines:
        ax1.legend(loc="upper right")

    ax2 = fig.add_subplot(gs[0, 1])
    ax2.axis("off")
    rows = [
        "Total time in bed, min",
        "Total sleep, min",
        "Sleep efficiency, %",
        "Sleep latency, min",
        "WASO, min",
        "Stage 1, %",
        "Stage 2, %",
        "SWS, %",
        "NREM, %",
        "REM, %",
    ]
    keys = ["TIB", "TST", "SE", "SOL", "WASO", "%N1", "%N2", "%N3", "%NREM", "%REM"]
    cell_text = [[stats_fmt[k]] for k in keys]

    if art_result is not None and art_result["n_total"] > 0:
        rows += ["Artifact epochs removed", "Artifact rejection, %"]
        cell_text += [
            [str(art_result["n_bad"])],
            [f"{(art_result['n_bad'] / max(art_result['n_total'], 1)) * 100:.2f}"],
        ]

    table = ax2.table(cellText=cell_text, rowLabels=rows, colWidths=(0.45, 0.25), loc="center")
    table.set_fontsize(12)
    table.scale(1.1, 1.9)

    fig.savefig(panel_png, bbox_inches="tight", dpi=150)
    plt.close(fig)

    return {
        "hypno_png": hypno_png,
        "stats_csv": stats_csv,
        "psd_stats_png": panel_png,
        "sleep_stats": stats_fmt,
        "artifact_csv": art_csv if use_yasa_art_detect else None,
        "artifact_rejection": art_result,
    }



def _stack_images_vertical(top_path: Path, bottom_path: Path, out_path: Path) -> Path:
    im_top = Image.open(top_path).convert("RGB")
    im_bottom = Image.open(bottom_path).convert("RGB")
    width = max(im_top.width, im_bottom.width)

    if im_top.width != width:
        im_top = im_top.resize((width, int(im_top.height * width / im_top.width)), Image.Resampling.LANCZOS)
    if im_bottom.width != width:
        im_bottom = im_bottom.resize((width, int(im_bottom.height * width / im_bottom.width)), Image.Resampling.LANCZOS)

    combined = Image.new("RGB", (width, im_top.height + im_bottom.height), (255, 255, 255))
    combined.paste(im_top, (0, 0))
    combined.paste(im_bottom, (0, im_top.height))
    combined.save(out_path)
    return out_path

def save_hypnospectrogram_fallback(
    eeg_path, outdir, hypnotxt, hypnofreq, eeg_channel, outpng, 
    winsec=10, freqrange=(0, 40), cmap="Spectral_r", overlap=True, 
    signal_name="EMG-R", caisr_nrem_events=None, **kwargs
):
    from sleepeegpy.pipeline import SpectralPipe
    import matplotlib.pyplot as plt
    import numpy as np
    import pandas as pd
    from pathlib import Path

    spectralpipe = SpectralPipe(
        path_to_eeg=str(eeg_path),
        output_dir=str(outdir),
        path_to_hypno=str(hypnotxt),
        hypno_freq=hypnofreq,
    )

    plt.close("all")
    spectralpipe.plot_hypnospectrogram(
        picks=eeg_channel, win_sec=winsec, freq_range=freqrange,
        cmap=cmap, overlap=overlap, save=False
    )

    fig = plt.gcf()
    
    events_path = Path(outdir) / "rbdtector_case" / "RBDtector output" / "RBDtection_Events_RBDtector output.csv"
    
    if events_path.exists() and len(fig.axes) > 0:
        try:
            ax_hypno = fig.axes[0]
            events_df = pd.read_csv(events_path, header=None, engine='python')
            
            base_zero = pd.Timestamp("1985-01-01 00:00:00")
            clean_sig = signal_name.replace("-", "").upper()
            
            hypno_data = np.loadtxt(hypnotxt)
            t_max_hours = len(hypno_data) / 3600.0

            found = 0
            for _, row in events_df.iterrows():
                if len(row) < 3: continue
                label = str(row[2]).replace("-","").upper()
                
                if clean_sig in label:
                    start_ev = pd.to_datetime(row[0])
                    end_ev = pd.to_datetime(row[1])
                    
                    ds = (start_ev - base_zero).total_seconds()
                    de = (end_ev - base_zero).total_seconds()
                    sh = ds / 3600.0
                    eh = de / 3600.0

                    idx = int(ds)
                    # ИСПРАВЛЕНИЕ: Рисуем RBD только если стадия = REM (4)
                    if idx >= 0 and idx < len(hypno_data) and hypno_data[idx] == 4:
                        if eh > 0 and sh < t_max_hours:
                            ax_hypno.hlines(y=-0.15, xmin=sh, xmax=eh, 
                                           color='red', linewidth=10, alpha=1.0, 
                                           transform=ax_hypno.get_xaxis_transform(),
                                           clip_on=False, zorder=100,
                                           label='Potential Rem Sleep Without Atonia' if found == 0 else "")
                            found += 1
            print(f"DEBUG: Successfully read CSV directly. Plotted {found} True REM RBD events.")
        except Exception as e:
            print(f"DEBUG ERROR during event plotting: {e}")

    # --- БЛОК ДЛЯ NREM ДВИЖЕНИЙ НОГ ---
    if caisr_nrem_events is not None and not caisr_nrem_events.empty and len(fig.axes) > 0:
        try:
            ax_hypno = fig.axes[0]
            plotted_nrem = 0
            hypno_data = np.loadtxt(hypnotxt)
            t_max_hours = len(hypno_data) / 3600.0

            for _, row in caisr_nrem_events.iterrows():
                sh = row['start_sec'] / 3600.0
                eh = row['end_sec'] / 3600.0
                
                display_eh = max(eh, sh + (30.0 / 3600.0))

                if eh > 0 and sh < t_max_hours:
                    ax_hypno.hlines(y=-0.30, xmin=sh, xmax=display_eh, 
                                   color='blue', linewidth=10, alpha=1.0, 
                                   transform=ax_hypno.get_xaxis_transform(),
                                   clip_on=False, zorder=100,
                                   label='Potential Periodic limb movement' if plotted_nrem == 0 else "")
                    plotted_nrem += 1
            print(f"DEBUG: Plotted {plotted_nrem} CAISR NREM events on hypnospectrogram.")
        except Exception as e:
            print(f"DEBUG ERROR during CAISR plotting: {e}")

    if len(fig.axes) > 0:
        ax_hypno = fig.axes[0]
        handles, labels = ax_hypno.get_legend_handles_labels()
        if handles:
            fig.legend(handles, labels, loc='upper center', bbox_to_anchor=(0.5, -0.25), ncol=2, fontsize=10, frameon=False)

    outpng.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(outpng, bbox_inches="tight", dpi=300)
    plt.close(fig)
    return outpng

def _save_hypnogram_with_rbdtector_events_png(
    path: Path,
    hypno_1hz_int: np.ndarray,
    calculated_data: pd.DataFrame,
    *,
    title: str | None = None,
    start_dt: pd.Timestamp | None = None,
    signal: str = "EMG-R",
    event_color: str = "#E63946",  # Modern red
    event_linewidth: float = 6.0,  
    event_y: float = -4.4,         
    verbose: bool = True,
) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    hypno_1hz_int = np.asarray(hypno_1hz_int, dtype=int).ravel()
    
    t_hours = np.arange(hypno_1hz_int.size) / 3600.0
    remapped = pd.Series(hypno_1hz_int).map({-2:-2, -1:-1, 0:0, 1:2, 2:3, 3:4, 4:1}).fillna(-2).to_numpy()
    y_vals = -1.0 * remapped

    # Use a clean, professional style
    plt.style.use("seaborn-v0_8-whitegrid")
    fig, ax = plt.subplots(figsize=(14, 5), dpi=300)
    
    # Background grid for readability
    ax.grid(axis='x', linestyle='--', alpha=0.5, color='#CCCCCC')
    ax.grid(axis='y', linestyle='-', alpha=0.3, color='#CCCCCC')

    # Draw the main hypnogram line with a softer black/grey
    ax.step(t_hours, y_vals, where="post", color="#333333", linewidth=1.5, alpha=0.8)

    masks = {
        "REM": np.ma.masked_not_equal(remapped, 1),
        "W": np.ma.masked_not_equal(remapped, 0),
        "N1": np.ma.masked_not_equal(remapped, 2),
        "N2": np.ma.masked_not_equal(remapped, 3),
        "N3": np.ma.masked_not_equal(remapped, 4),
    }
    
    # Modern color palette for sleep stages
    colors = {
        "REM": "#E63946",  # Red
        "W": "#F4A261",    # Orange/Sand
        "N1": "#A8DADC",   # Light Blue
        "N2": "#457B9D",   # Medium Blue
        "N3": "#1D3557"    # Dark Blue
    }
    
    for stage, arr in masks.items():
        ax.step(t_hours, -1.0 * arr, where="post", color=colors[stage], linewidth=2.5)
        
        # Add shading under REM sleep to make it stand out
        if stage == "REM":
            ax.fill_between(t_hours, -1.0 * arr, 0, step="post", color=colors[stage], alpha=0.2)

    if calculated_data is not None and not calculated_data.empty:
        try:
            clean_sig = signal.replace("-", "").upper()
            rec_start = start_dt if start_dt else pd.to_datetime(calculated_data.iloc[0, 0])
            
            def to_day_seconds(dt):
                if isinstance(dt, str): dt = pd.to_datetime(dt)
                return dt.hour * 3600 + dt.minute * 60 + dt.second + dt.microsecond / 1e6

            rec_start_sec = to_day_seconds(rec_start)
            plotted_rbd = 0
            
            for _, row in calculated_data.iterrows():
                if len(row) < 3: continue
                label = str(row[2]).replace("-","").upper()
                
                if clean_sig in label:
                    ev_start_sec = to_day_seconds(row[0])
                    ev_end_sec = to_day_seconds(row[1])

                    diff_start = ev_start_sec - rec_start_sec
                    diff_end = ev_end_sec - rec_start_sec

                    if diff_start < -43200: diff_start += 86400
                    if diff_end < -43200: diff_end += 86400
                    
                    idx = int(diff_start)
                    if idx >= 0 and idx < len(hypno_1hz_int) and hypno_1hz_int[idx] == 4:
                        start_h, end_h = diff_start / 3600.0, diff_end / 3600.0
                        if end_h > 0 and start_h < t_hours[-1]:
                            ax.hlines(y=event_y, xmin=max(0, start_h), xmax=min(t_hours[-1], end_h),
                                      color=event_color, linewidth=event_linewidth, alpha=1.0, 
                                      label='Potential RSWA' if plotted_rbd == 0 else "")
                            plotted_rbd += 1
            
            if verbose: print(f"Plotted {plotted_rbd} True REM events for {signal} at y={event_y}")
        except Exception as e:
            if verbose: print(f"Event Plotting Error: {e}")

    ax.set_yticks([0, -1, -2, -3, -4])
    ax.set_yticklabels(["W", "REM", "N1", "N2", "N3"], fontweight='bold', color='#444444')
    ax.set_ylim(-4.8, 0.5) 
    ax.set_xlim(0, max(t_hours[-1] if t_hours.size else 0, 0.01))
    ax.set_xlabel("Time [hrs]", fontweight='bold', color='#444444')
    
    if title: 
        ax.set_title(title, pad=15, fontweight='bold', fontsize=14, color='#222222')
    
    handles, labels = ax.get_legend_handles_labels()
    if handles: 
        ax.legend(handles, labels, loc='lower center', bbox_to_anchor=(0.5, -0.2), ncol=2, frameon=False)

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color('#888888')
    ax.spines["bottom"].set_color('#888888')
    
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)

def _make_psd_stats_panel(
    *,
    raw: mne.io.BaseRaw,
    hypno_1hz_int: np.ndarray,
    subject_id: str,
    outdir: Path,
    eeg_channel_for_report: str,
    separate_nrem_power_spectrum: bool = True,
    use_yasa_art_detect: bool = True,
    yasa_art_epoch_sec: int = 30,
    yasa_art_window: int = 5,
    yasa_art_method: str = "covar",
    yasa_art_threshold: float = 3.0,
    yasa_art_include: tuple[int, ...] = (0, 1, 2, 3, 4),
    verbose: bool = True,
) -> dict[str, Any]:
    outdir.mkdir(parents=True, exist_ok=True)
    raw_one = raw.copy().pick([eeg_channel_for_report])
    _ensure_loaded(raw_one)

    hypno_1hz_int = np.asarray(hypno_1hz_int, dtype=int)
    stats = _compute_sleep_stats(hypno_1hz_int)
    stats_fmt = _format_sleep_stats(stats)

    hypno_png = outdir / f"{subject_id}_hypno.png"
    stats_csv = outdir / f"{subject_id}_stats.csv"
    panel_png = outdir / f"{subject_id}_psd_stats.png"
    art_csv = outdir / f"{subject_id}_yasa_artifacts.csv"

    _save_hypnogram_png(hypno_png, hypno_1hz_int, title=subject_id)
    _save_kv_csv(stats_csv, stats_fmt)

    art_result = None
    epoch_keep_mask = None

    if use_yasa_art_detect:
        art_result = _run_yasa_art_detect(
            raw_one,
            hypno_1hz_int,
            epoch_sec=yasa_art_epoch_sec,
            include=tuple(yasa_art_include),
            window=yasa_art_window,
            method=yasa_art_method,
            threshold=yasa_art_threshold,
            verbose=verbose,
        )
        epoch_keep_mask = art_result["epoch_mask_keep"]

        _save_artifact_epochs_csv(
            art_csv,
            hypno_30=art_result["hypno_30"],
            bad_mask=art_result["epoch_mask_bad"],
            zscores=art_result["zscores"],
        )

    epochs, _ = _make_stage_epochs(
        raw_one,
        hypno_1hz_int,
        epoch_sec=30,
        epoch_keep_mask=epoch_keep_mask,
    )

    stage_order, stage_psd = _compute_stage_psd(
        epochs,
        separate_nrem_power_spectrum=separate_nrem_power_spectrum,
        fmin=0.5,
        fmax=40.0,
    )

    plt.style.use("seaborn-v0_8-whitegrid")
    fig = plt.figure(constrained_layout=True, figsize=(14, 6), dpi=300)
    gs = fig.add_gridspec(1, 2, width_ratios=[1.4, 1.0])

    ax1 = fig.add_subplot(gs[0, 0])
    
    # Professional colors for PSD matching the hypnogram
    stage_color = {"W": "#F4A261", "N2": "#457B9D", "N3": "#1D3557", "REM": "#E63946", "NREM": "#1D3557"}
    stage_label = {"W": "Wake", "N2": "N2", "N3": "N3", "REM": "REM", "NREM": "NREM"}

    for stage in stage_order:
        if stage in stage_psd:
            freqs, psd_db = stage_psd[stage]
            ax1.plot(freqs, psd_db, color=stage_color[stage], label=stage_label[stage], linewidth=2.5)
            # Add subtle fill underneath the curves
            ax1.fill_between(freqs, psd_db, np.min(psd_db) - 5, color=stage_color[stage], alpha=0.1)

    ax1.set_title(f"Power Spectral Density by Stage - {subject_id}", pad=15, fontweight='bold', fontsize=14)
    ax1.set_ylabel("Power Spectral Density (dB)", fontweight='bold')
    ax1.set_xlabel("Frequency (Hz)", fontweight='bold')
    ax1.set_xlim(freqs.min(), freqs.max())
    
    # Hide top/right spines for cleaner look
    ax1.spines["top"].set_visible(False)
    ax1.spines["right"].set_visible(False)
    
    if ax1.lines:
        ax1.legend(loc="upper right", frameon=True, shadow=False, edgecolor='#CCCCCC')

    ax2 = fig.add_subplot(gs[0, 1])
    ax2.axis("off")
    rows = [
        "Total time in bed (min)",
        "Total sleep (min)",
        "Sleep efficiency (%)",
        "Sleep latency (min)",
        "WASO (min)",
        "Stage 1 (%)",
        "Stage 2 (%)",
        "SWS (%)",
        "NREM (%)",
        "REM (%)",
    ]
    keys = ["TIB", "TST", "SE", "SOL", "WASO", "%N1", "%N2", "%N3", "%NREM", "%REM"]
    cell_text = [[stats_fmt[k]] for k in keys]

    if art_result is not None and art_result["n_total"] > 0:
        rows += ["Artifact epochs removed", "Artifact rejection (%)"]
        cell_text += [
            [str(art_result["n_bad"])],
            [f"{(art_result['n_bad'] / max(art_result['n_total'], 1)) * 100:.2f}"],
        ]

    # Create a cleaner table without heavy borders
    table = ax2.table(cellText=cell_text, rowLabels=rows, colWidths=(0.3, 0.2), loc="center")
    table.auto_set_font_size(False)
    table.set_fontsize(11)
    table.scale(1.2, 2.0)
    
    # Clean up table cell edges
    for (row, col), cell in table.get_celld().items():
        cell.set_edgecolor('#E0E0E0')
        if col == -1:  # Row labels
            cell.set_text_props(weight='bold', color='#333333')
            cell.set_facecolor('#F8F9FA')

    fig.savefig(panel_png, bbox_inches="tight")
    plt.close(fig)

    return {
        "hypno_png": hypno_png,
        "stats_csv": stats_csv,
        "psd_stats_png": panel_png,
        "sleep_stats": stats_fmt,
        "artifact_csv": art_csv if use_yasa_art_detect else None,
        "artifact_rejection": art_result,
    }

def export_events_to_pdf(
    raw: mne.io.BaseRaw,
    out_path: Path,
    outdir: Path, 
    caisr_nrem_events: pd.DataFrame,
    start_dt: pd.Timestamp,
    rbd_channel: str,
    rbd_match_label: str,
    plm_l_channel: str,
    plm_r_channel: str | None, 
    eeg_channels: list[str], 
    eog_channels: list[str],   
    hypno_1hz_int: np.ndarray,
    epoch_sec: float = 30.0
):
    path = Path(out_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    
    print(f"\n[REPORT] Generating Events Atlas PDF at {path.name}...")
    from matplotlib.backends.backend_pdf import PdfPages
    
    with PdfPages(path) as pdf:
        # --- 1. PLM EVENTS (BLUE) ---
        if caisr_nrem_events is not None and not caisr_nrem_events.empty:
            plm_epochs = {}
            for idx, row in caisr_nrem_events.iterrows():
                sh = row['start_sec']
                eh = row['end_sec']
                s_ep = int(sh // epoch_sec)
                e_ep = int(eh // epoch_sec)
                for ep in range(s_ep, e_ep + 1):
                    if ep not in plm_epochs:
                        plm_epochs[ep] = []
                    plm_epochs[ep].append((sh, eh))
            
            print(f"[REPORT] Plotting {min(len(plm_epochs), MAX_ATLAS_PAGES_PER_TYPE)} of {len(plm_epochs)} 30-s epochs with PLM events...")
            # ИСПРАВЛЕНИЕ ПОРЯДКА: Сначала ЭОГ, потом ЭЭГ, потом ЭМГ ног
            plm_channels_to_plot = eog_channels + eeg_channels + [ch for ch in [plm_l_channel, plm_r_channel] if ch]
            
            for _page_i, ep in enumerate(sorted(plm_epochs.keys())):
                if _page_i >= MAX_ATLAS_PAGES_PER_TYPE:
                    print(f"[REPORT] PLM atlas truncated at {MAX_ATLAS_PAGES_PER_TYPE} pages.")
                    break
                ep_start = ep * epoch_sec
                ep_end = ep_start + epoch_sec
                
                abs_start_str = (start_dt + pd.Timedelta(seconds=ep_start)).strftime('%H:%M:%S')
                abs_end_str = (start_dt + pd.Timedelta(seconds=ep_end)).strftime('%H:%M:%S')
                title = f"PLM Events (Epoch {ep+1}: {abs_start_str} - {abs_end_str})"
                
                _plot_30s_epoch(pdf, raw, ep_start, ep_end, plm_channels_to_plot, plm_epochs[ep], title, 'blue', start_dt)
                
        # --- 2. RBD EVENTS (RED) ---
        events_path = Path(outdir) / "rbdtector_case" / "RBDtector output" / "RBDtection_Events_RBDtector output.csv"
        if events_path.exists():
            try:
                rbd_df = pd.read_csv(events_path, header=None, engine='python')
                base_zero = pd.Timestamp("1985-01-01 00:00:00")
                clean_sig = rbd_match_label.replace("-", "").upper()
                
                rbd_epochs = {}
                for _, row in rbd_df.iterrows():
                    if len(row) < 3: continue
                    label = str(row[2]).replace("-", "").upper()
                    
                    if clean_sig in label:
                        start_ev = pd.to_datetime(row[0])
                        end_ev = pd.to_datetime(row[1])
                        
                        ds = (start_ev - base_zero).total_seconds()
                        de = (end_ev - base_zero).total_seconds()
                        
                        idx = int(ds)
                        if idx >= 0 and idx < len(hypno_1hz_int) and hypno_1hz_int[idx] == 4:
                            s_ep = int(ds // epoch_sec)
                            e_ep = int(de // epoch_sec)
                            for ep in range(s_ep, e_ep + 1):
                                if ep not in rbd_epochs:
                                    rbd_epochs[ep] = []
                                rbd_epochs[ep].append((ds, de))
                
                print(f"[REPORT] Plotting {min(len(rbd_epochs), MAX_ATLAS_PAGES_PER_TYPE)} of {len(rbd_epochs)} 30-s epochs with True REM events...")
                # ИСПРАВЛЕНИЕ ПОРЯДКА: Сначала ЭОГ, потом ЭЭГ, потом ЭМГ подбородка
                rbd_channels_to_plot = eog_channels + eeg_channels + [rbd_channel]
                
                for _page_i, ep in enumerate(sorted(rbd_epochs.keys())):
                    if _page_i >= MAX_ATLAS_PAGES_PER_TYPE:
                        print(f"[REPORT] RBD atlas truncated at {MAX_ATLAS_PAGES_PER_TYPE} pages.")
                        break
                    ep_start = ep * epoch_sec
                    ep_end = ep_start + epoch_sec
                    
                    abs_start_str = (start_dt + pd.Timedelta(seconds=ep_start)).strftime('%H:%M:%S')
                    abs_end_str = (start_dt + pd.Timedelta(seconds=ep_end)).strftime('%H:%M:%S')
                    title = f"RBD Events (Epoch {ep+1}: {abs_start_str} - {abs_end_str})"
                    
                    _plot_30s_epoch(pdf, raw, ep_start, ep_end, rbd_channels_to_plot, rbd_epochs[ep], title, 'red', start_dt)
            except Exception as e:
                print(f"[REPORT] Failed to plot RBD events: {e}")

    print("[REPORT] Events Atlas successfully saved!")

def _plot_30s_epoch(pdf, raw, start_sec, end_sec, channels_to_plot, events, title, color, start_dt):
    channels_to_plot = [ch for ch in channels_to_plot if ch in raw.ch_names]
    if not channels_to_plot:
        return
        
    try:
        data, times = raw.copy().pick(channels_to_plot).get_data(tmin=start_sec, tmax=end_sec, return_times=True)
    except Exception as e:
        print(f"Failed to extract data: {e}")
        return
        
    n_plots = len(channels_to_plot)
    fig, axes = plt.subplots(n_plots, 1, figsize=(14, 1.8 * n_plots), sharex=True, gridspec_kw={'hspace': 0.1})
    if n_plots == 1:
        axes = [axes]
        
    for i, ch_name in enumerate(channels_to_plot):
        ch_upper = ch_name.upper()
        
        # Индивидуальные масштабы
        if 'EOG' in ch_upper:
            y_limit = 200.0  
            sig_color = "#88bce0" 
        elif 'EEG' in ch_upper or ch_upper in ['F', 'C', 'O']:
            y_limit = 80.0   
            sig_color = '#2c3e50' 
        else:
            y_limit = 50.0   
            sig_color = "#e07f7f" 
            
        axes[i].plot(times, data[i] * 1e6, color=sig_color, linewidth=0.8)
        
        axes[i].set_ylim(-y_limit, y_limit)
        axes[i].set_ylabel(f"{ch_name}\n(µV)", fontsize=9)
        
        for (sh, eh) in events:
            h_start = max(start_sec, sh)
            h_end = min(end_sec, eh)
            if h_end > h_start:
                axes[i].axvspan(h_start, h_end, color=color, alpha=0.25)
                
        axes[i].grid(True, axis='x', linestyle='--', alpha=0.5)
        axes[i].spines['top'].set_visible(False)
        axes[i].spines['right'].set_visible(False)
        
    axes[0].set_title(title, fontsize=12, fontweight='bold', pad=10)
    axes[-1].set_xlim(start_sec, end_sec)
    
    ticks = axes[-1].get_xticks()
    labels = []
    for t in ticks:
        abs_time = start_dt + pd.Timedelta(seconds=t)
        labels.append(abs_time.strftime('%H:%M:%S'))
        
    axes[-1].set_xticks(ticks)
    axes[-1].set_xticklabels(labels)
    axes[-1].set_xlabel("Absolute Time", fontsize=11)
    
    pdf.savefig(fig, bbox_inches='tight')
    plt.close(fig)


def cleanup_output_directory(outdir: Path, verbose: bool = True):
    """Переносит весь технический мусор во внутреннюю папку Analysis_Details."""
    try:
        details_dir = outdir / "Analysis_Details"
        details_dir.mkdir(exist_ok=True)
        
        # Список расширений, которые мы ХОТИМ оставить в корне
        keep_extensions = {'.pdf', '.fif', '.docx'}
        
        for item in outdir.iterdir():
            if item.is_file():
                # Переносим всё, что НЕ входит в список исключений
                if item.suffix.lower() not in keep_extensions:
                    shutil.move(str(item), str(details_dir / item.name))
                    
        if verbose:
            print(f"[REPORT] Cleanup complete. Intermediate files moved to {details_dir.name}/")
    except Exception as e:
        print(f"[REPORT] Failed to organize output folder: {e}")
def mff_yasa_sleepeegpy_combo(
    mff_paths: list[str | Path] | str | Path,
    *,
    montage_type: str = "adult_e256",
    side: str = "right",
    site: str = "C",
    show_preview: bool = False,
    outdir: str | Path | None = None,
    outdir_suffix: str = "_spectral",
    hypno_freq: int = 1,
    hypno_name: str = "hypno_point_per_row_1hz_numeric.txt",
    fix_epochs_endtime: bool = True,
    step: int = 1000,
    max_k: int = 20000,
    attempts_dirname: str = "_epochs_attempts",
    delete_events_user_markup: bool = False,
    read_raw_kwargs: Optional[dict] = None,
    preload: bool = True,
    mne_verbose: str | None = "ERROR",
    win_sec: float = 10,
    freq_range: Tuple[float, float] = (0, 40),
    cmap: str = "Spectral_r",
    overlap: bool = True,
    pie_name: str = "pie.png",
    hypnospec_name: str = "hypnospec.png",
    combo_name: str = "combo.png",
    fif_name: str = "sleep_6ch_ref_filter_raw.fif",
    run_rbdtector_headless: bool = False,
    make_sleep_report: bool = True,
    report_name: str = "sleep_report.png",
    make_docx_report: bool = False,
    subject_age: int | None = None,
    sex: str | None = None,
    spo2_edf_path: str | Path | None = None,
    spo2_channel: str = "SaO2 SPO2",
    pleth_channel: str = "Pulse Pleth",
    apply_prefilter: bool = True,
    prefilter_l_freq: float = 0.1,
    prefilter_h_freq: float = 70.0,
    notch_base_freq: float = 50.0,
    use_yasa_art_detect: bool = True,
    yasa_art_epoch_sec: int = 30,
    yasa_art_window: int = 5,
    yasa_art_method: str = "covar",
    yasa_art_threshold: float = 3.0,
    yasa_art_include: tuple[int, ...] = (0, 1, 2, 3, 4),
    verbose: bool = True,
    subject_name: str | None = None,
    append_pdf_path: str | Path | None = None,
    primary_scoring: str = "YASA",
    anchor_dir: str | Path | None = None,
    scoring_check: bool = False,
) -> Dict[str, Any]:

    _timer = PipelineTimer()

    # 1. Нормализуем mff_paths в список
    if isinstance(mff_paths, (str, Path)):
        mff_paths = [mff_paths]
    mff_paths = [Path(p) for p in mff_paths]
    
    if not mff_paths:
        raise ValueError("No MFF paths provided.")

    mff_path_primary = mff_paths[0] # Используем первый файл для именования всех отчетов
    read_raw_kwargs = {} if read_raw_kwargs is None else dict(read_raw_kwargs)

    montage_type = montage_type.lower().strip()
    side = side.lower().strip()
    site = site.upper().strip()

    if montage_type not in CFGS:
        raise ValueError(f"montage_type must be one of {list(CFGS.keys())}, got: {montage_type}")

    cfg_set = CFGS[montage_type]

    if side not in cfg_set:
        raise ValueError(f"side must be one of {list(cfg_set.keys())}, got: {side}")
    if site not in ("F", "C", "O"):
        raise ValueError("site must be 'F', 'C', or 'O'")

    stem = mff_path_primary.name[:-4] if mff_path_primary.name.lower().endswith(".mff") else mff_path_primary.name
    subject_id = stem.split('_')[0]

    if outdir is None:
        if anchor_dir:
            outdir = Path(anchor_dir) / "SleepReport" / subject_id
        else:
            outdir = mff_path_primary.parent / f"{stem}{outdir_suffix}"
    else:
        outdir = Path(outdir)
        
    outdir.mkdir(parents=True, exist_ok=True)
    
    details_dir = outdir / "Analysis_Details"
    details_dir.mkdir(parents=True, exist_ok=True)

    c = cfg_set[side]
    eeg_F = c["F"]
    eeg_C = c["C"]
    eeg_O = c["O"]
    mastoid = c["M"]
    eog1_src = c["EOG1"]
    eog2_src = c["EOG2"]
    emg_src = c["EMG"]

    eeg_alias_F = c["EEG_ALIAS"]["F"]
    eeg_alias_C = c["EEG_ALIAS"]["C"]
    eeg_alias_O = c["EEG_ALIAS"]["O"]
    eog1_alias = c["EOG1_ALIAS"]
    eog2_alias = c["EOG2_ALIAS"]
    emg_alias = c["EMG_ALIAS"]

    eeg_alias_for_yasa = c["EEG_ALIAS"][site]
    eeg_alias_for_report = c["EEG_ALIAS"]["C"]

    LM_src = cfg_set["left"]["M"]
    RM_src = cfg_set["right"]["M"]

    hypno_txt = details_dir / hypno_name
    pie_png = details_dir / pie_name
    hspec_png = details_dir / hypnospec_name
    combo_png = details_dir / combo_name
    fif_path = details_dir / fif_name
    report_png = details_dir / report_name
    artifact_csv = details_dir / f"{mff_path_primary.stem}_yasa_artifacts.csv"

    preview_png = details_dir / "preview.png"
    psd_stats_png = details_dir / "psd_stats.png"
    docx_report_path = outdir / f"sleep_report_{mff_path_primary.stem}.docx"
    legacy_pie_png = details_dir / f"{mff_path_primary.stem}_sleep_stages.png"
    hypno_rbd_events_png = details_dir / f"{mff_path_primary.stem}_hypno_rbd_events.png"
    caisr_motor_events_png = outdir / f"{mff_path_primary.stem}_motor_events.png"

    need = [eeg_F, eeg_C, eeg_O, mastoid, eog1_src, eog2_src, emg_src]
    need_ref = [eog1_src, eog2_src, eeg_F, eeg_C, eeg_O, emg_src, LM_src, RM_src]

    if scoring_check:
        for s in ["left", "right"]:
            c_s = cfg_set[s]
            for ch in [c_s["F"], c_s["C"], c_s["O"], c_s["M"], c_s["EOG1"], c_s["EOG2"], c_s["EMG"]]:
                if ch not in need: need.append(ch)
                if ch not in need_ref: need_ref.append(ch)

    def _process_single_mff(single_mff_path: Path):
        """Helper to load, fix epochs, and extract raw + raw_ref for a single MFF file."""
        def _read_picked(picks_to_load):
            r = safe_read_mff(str(single_mff_path), preload=False, verbose=mne_verbose, **read_raw_kwargs)
            caisr_extras = [ch for ch in r.ch_names if any(x in ch.upper() for x in ['ECG', 'LEG', 'CHEST', 'ABD', 'SPO2', 'SAO2'])]
            for ch in caisr_extras:
                if ch not in picks_to_load: picks_to_load.append(ch)
            missing = [ch for ch in picks_to_load if ch not in r.ch_names]
            if missing: raise RuntimeError(f"Missing channels in {single_mff_path.name}: {missing}")
            
            try:
                r.pick(picks_to_load)
                r.load_data()
            except Exception:
                r.load_data()
                r.pick(picks_to_load)
                
            if apply_prefilter:
                r = _apply_mff_prefilter(
                    r, picks=picks_to_load, l_freq=prefilter_l_freq, h_freq=prefilter_h_freq, 
                    notch_base=notch_base_freq, verbose=verbose
                )
            return r

        if fix_epochs_endtime:
            epochs_xml = single_mff_path / "epochs.xml"
            if not epochs_xml.exists(): 
                print(f"epochs.xml missing for {single_mff_path.name}, creating exactly...")
                fix_mff_epochs_xml(single_mff_path)
            
            attempts_dir = single_mff_path.parent / f"{single_mff_path.name}{attempts_dirname}"
            attempts_dir.mkdir(parents=True, exist_ok=True)
            
            markup = single_mff_path / "Events_User Markup.xml"
            if markup.exists():
                shutil.copy2(markup, attempts_dir / "Events_User Markup.xml")
                
            backup_xml = single_mff_path.parent / f"{single_mff_path.name}__epochs_original.xml"
            if not backup_xml.exists():
                shutil.copy2(epochs_xml, backup_xml)

            try:
                r_bip = _read_picked(need.copy())
                if verbose: print(f"SUCCESS: no shift needed for {single_mff_path.name}")
            except Exception as e:
                # Если ошибка связана с отсутствующими каналами, не пытаемся патчить epochs.xml 40000 раз
                if isinstance(e, RuntimeError) and "Missing channels" in str(e):
                    raise e
                
                original_text = backup_xml.read_text(encoding="utf-8", errors="replace")
                matches = list(_ENDTIME_RE.finditer(original_text))
                if not matches: raise RuntimeError("No <endTime> found in epochs.xml")
                
                last = matches[-1]
                orig_end = int(last.group(2))
                deltas = []
                for k in range(1, max_k + 1):
                    deltas.append(+k * step)
                    deltas.append(-k * step)
                    
                def make_patched_text(new_end):
                    s, e2 = last.span(2)
                    return original_text[:s] + str(new_end) + original_text[e2:]
                    
                r_bip = None
                for i, d in enumerate(deltas):
                    if i % 500 == 0:
                        print(f"[Recovery] Testing delta {d} (Attempt {i}/{len(deltas)})...")
                        
                    new_end = orig_end + d
                    epochs_xml.write_text(make_patched_text(new_end), encoding="utf-8")
                    try:
                        r_bip = _read_picked(need.copy())
                        print(f"[Recovery SUCCESS] Found correct endTime! Delta={d} worked.")
                        break
                    except Exception as inner_e:
                        if isinstance(inner_e, RuntimeError) and "Missing channels" in str(inner_e):
                            raise inner_e
                        pass
                if r_bip is None:
                    shutil.copy2(backup_xml, epochs_xml)
                    raise RuntimeError(f"No delta worked for {single_mff_path.name}. Original error: {e}")
        else:
            r_bip = _read_picked(need.copy())

        # Load reference raw since epochs.xml is now stable
        r_ref = _read_picked(need_ref.copy())
        return r_bip, r_ref

    # 2. Iterate through all MFF parts, fix and load them
    _timer.start("Load & fix MFF data")
    raw_parts = []
    raw_ref_parts = []
    for mff_p in mff_paths:
        if verbose: print(f"\n--- Loading and checking part: {mff_p.name} ---")
        rbip, rref = _process_single_mff(mff_p)
        raw_parts.append(rbip)
        raw_ref_parts.append(rref)
    _timer.stop()

    # 3. Merge parts together into one continuous recording
    if len(raw_parts) > 1:
        if verbose: print(f"\n>>> Merging {len(raw_parts)} parts into a single continuous record...")
        raw = mne.concatenate_raws(raw_parts)
        raw_ref = mne.concatenate_raws(raw_ref_parts)
    else:
        raw = raw_parts[0]
        raw_ref = raw_ref_parts[0]

    # --- Continue pipeline normally on the merged `raw` and `raw_ref` ---
    _timer.start("Bipolar referencing & channel export")
    r = raw.copy()
    r = mne.set_bipolar_reference(r, anode=eeg_F, cathode=mastoid, ch_name=eeg_alias_F, drop_refs=False, copy=True)
    r = mne.set_bipolar_reference(r, anode=eeg_C, cathode=mastoid, ch_name=eeg_alias_C, drop_refs=False, copy=True)
    r = mne.set_bipolar_reference(r, anode=eeg_O, cathode=mastoid, ch_name=eeg_alias_O, drop_refs=False, copy=True)
    r.set_channel_types({eeg_alias_F: "eeg", eeg_alias_C: "eeg", eeg_alias_O: "eeg"})
    r = mne.set_bipolar_reference(r, anode=eog1_src, cathode=mastoid, ch_name=eog1_alias, drop_refs=False, copy=True)
    r = mne.set_bipolar_reference(r, anode=eog2_src, cathode=mastoid, ch_name=eog2_alias, drop_refs=False, copy=True)
    r.set_channel_types({eog1_alias: "eog", eog2_alias: "eog"})
    r.rename_channels({emg_src: emg_alias})
    r.set_channel_types({emg_alias: "emg"})

    raw_ref.rename_channels({eog1_src: eog1_alias, eog2_src: eog2_alias, eeg_F: eeg_alias_F, eeg_C: eeg_alias_C, eeg_O: eeg_alias_O, emg_src: emg_alias, LM_src: "LM", RM_src: "RM"})
    raw_ref.set_channel_types({eog1_alias: "ecog", eog2_alias: "ecog", eeg_alias_F: "eeg", eeg_alias_C: "eeg", eeg_alias_O: "eeg", emg_alias: "emg", "LM": "eeg", "RM": "eeg"})
    raw_ref.set_eeg_reference(ref_channels=["RM", "LM"], ch_type="ecog")
    eeg_ref = "RM" if side == "left" else "LM"
    raw_ref.set_eeg_reference(ref_channels=[eeg_ref], ch_type="eeg")
    raw_ref.drop_channels(["LM", "RM"])
    raw_ref.set_channel_types({eog1_alias: "eog", eog2_alias: "eog"})
    emg_h_freq = min(100.0, prefilter_h_freq) if apply_prefilter else 100.0
    raw_ref.filter(10, emg_h_freq, picks=[emg_alias])

    export_order = [eeg_alias_F, eeg_alias_C, eeg_alias_O, eog1_alias, eog2_alias, emg_alias]
    raw_export = raw_ref.copy().pick(export_order)
    raw_export.reorder_channels(export_order)
    if verbose:
        print("Exporting channels:", len(raw_export.ch_names), raw_export.ch_names)
    assert len(raw_export.ch_names) == 6, raw_export.ch_names
    raw_export.save(str(fif_path), overwrite=True)

    _timer.stop()

    # ==========================================
    # RUN YASA FIRST (To generate stages_csv for CAISR)
    # ==========================================
    _timer.start("YASA sleep staging")
    
    hypno_int_30s = None
    raw_yasa = None
    hyp = None

    if scoring_check:
        if verbose:
            print(">>> Scoring Check mode enabled. Running YASA for all channels...")
        
        ch_names = {
            ("left", "F"): "F3", ("left", "C"): "C3", ("left", "O"): "O1",
            ("right", "F"): "F4", ("right", "C"): "C4", ("right", "O"): "O2"
        }
        
        for t_side in ["left", "right"]:
            other_side = "right" if t_side == "left" else "left"
            for t_site in ["F", "C", "O"]:
                c_test = CFGS[montage_type][t_side]
                c_other = CFGS[montage_type][other_side]
                
                t_anode = c_test[t_site]
                
                # Determine standard 10-20 names for fallbacks
                if t_site == "F":
                    std_site = "F3" if t_side == "left" else "F4"
                elif t_site == "C":
                    std_site = "C3" if t_side == "left" else "C4"
                else: # O
                    std_site = "O1" if t_side == "left" else "O2"
                std_mastoid = "M1" if t_side == "left" else "M2"
                
                def find_ch(preferred, alternatives):
                    if preferred in raw.ch_names: return preferred
                    for alt in alternatives:
                        if alt in raw.ch_names: return alt
                        # also check uppercase / lowercase / space variations
                        for ch in raw.ch_names:
                            if alt.lower() in ch.lower():
                                return ch
                    return preferred # return preferred so it shows up in missing if not found
                
                # Fallbacks
                t_anode = find_ch(c_test[t_site], [std_site])
                t_mastoid = find_ch(c_test["M"], [std_mastoid, "A1" if t_side == "left" else "A2"])
                if t_mastoid not in raw.ch_names:
                    t_mastoid = find_ch(c_other["M"], ["M2" if t_side == "left" else "M1", "A2" if t_side == "left" else "A1"])
                
                t_eog1 = find_ch(c_test["EOG1"], ["EOG", "LOC" if t_side == "left" else "ROC", "E1" if t_side == "left" else "E2"])
                if t_eog1 not in raw.ch_names:
                    t_eog1 = find_ch(c_other["EOG1"], ["EOG", "ROC" if t_side == "left" else "LOC", "E2" if t_side == "left" else "E1"])
                
                t_emg = find_ch(c_test["EMG"], ["EMG", "CHIN"])
                if t_emg not in raw.ch_names:
                    t_emg = find_ch(c_other["EMG"], ["EMG", "CHIN"])
                
                t_alias_eeg = c_test["EEG_ALIAS"][t_site]
                t_alias_eog1 = c_test["EOG1_ALIAS"]
                t_alias_emg = c_test["EMG_ALIAS"]
                
                req_chans = [t_anode, t_mastoid, t_eog1, t_emg]
                missing = [ch for ch in req_chans if ch not in raw.ch_names]
                
                if not missing:
                    try:
                        r_test = raw.copy()
                        r_test = mne.set_bipolar_reference(r_test, anode=t_anode, cathode=t_mastoid, ch_name=t_alias_eeg, drop_refs=False, copy=True, verbose=False)
                        r_test.set_channel_types({t_alias_eeg: "eeg"}, verbose=False)
                        
                        r_test = mne.set_bipolar_reference(r_test, anode=t_eog1, cathode=t_mastoid, ch_name=t_alias_eog1, drop_refs=False, copy=True, verbose=False)
                        r_test.set_channel_types({t_alias_eog1: "eog"}, verbose=False)
                        
                        r_test.rename_channels({t_emg: t_alias_emg})
                        r_test.set_channel_types({t_alias_emg: "emg"}, verbose=False)
                        
                        raw_yasa_test = r_test.pick([t_alias_eeg, t_alias_eog1, t_alias_emg])
                        sls_test = yasa.SleepStaging(raw_yasa_test, eeg_name=t_alias_eeg, eog_name=t_alias_eog1, emg_name=t_alias_emg)
                        hyp_test = sls_test.predict()
                        
                        h_test = hyp_test.hypno if hasattr(hyp_test, "hypno") else hyp_test
                        h_test = np.asarray(h_test)
                        
                        if not np.issubdtype(h_test.dtype, np.integer):
                            h_test = yasa.hypno_str_to_int(h_test.astype(str))
                        else:
                            h_test = h_test.astype(int)
                            
                        channel_label = ch_names[(t_side, t_site)]
                        txt_path = outdir / f"hypno_yasa_{channel_label}.txt"
                        pd.Series(h_test).to_csv(txt_path, index=False, header=False)
                        
                        if verbose:
                            print(f">>> Generated {txt_path.name}")
                            
                        if hypno_int_30s is None:
                            hypno_int_30s = h_test
                            raw_yasa = raw_yasa_test
                            hyp = hyp_test
                            if verbose:
                                print(f">>> Using {channel_label} as the primary scoring for the pipeline.")
                    except Exception as e:
                        if verbose:
                            print(f">>> Failed to process channel {ch_names[(t_side, t_site)]}: {e}")
                else:
                    if verbose:
                        print(f">>> [WARNING] Skipping {ch_names[(t_side, t_site)]} due to missing channels in raw data: {missing}")

    if hypno_int_30s is None:
        raw_yasa = r.copy().pick([eeg_alias_for_yasa, eog1_alias, emg_alias])
        sls = yasa.SleepStaging(raw_yasa, eeg_name=eeg_alias_for_yasa, eog_name=eog1_alias, emg_name=emg_alias)
        hyp = sls.predict()
        hypno = hyp.hypno if hasattr(hyp, "hypno") else hyp
        hypno = np.asarray(hypno)
    
        if np.issubdtype(hypno.dtype, np.integer):
            hypno_int_30s = hypno.astype(int)
        else:
            hypno_str_30s = hypno.astype(str)
            hypno_int_30s = yasa.hypno_str_to_int(hypno_str_30s)

    forced_stages_csv = details_dir / "stages.csv"
    pd.DataFrame({"stage": hypno_int_30s}).to_csv(forced_stages_csv, index=False)
    if verbose:
        print(f">>> YASA stages saved to {forced_stages_csv} to feed CAISR limb module.")

    _timer.stop()

    # ==========================================
    # RUN CAISR PIPELINE
    # ==========================================
    _timer.start("CAISR pipeline (staging + limb)")
    caisr_results = {}
    try:
        if verbose:
            print(">>> Running full CAISR pipeline...")
        caisr_results = run_full_caisr_pipeline(
            raw,
            subject_id=mff_path_primary.stem,
            outdir=outdir,
            stages_csv_path=str(forced_stages_csv)
        )
    except Exception as e:
        print("\n" + "="*60)
        print(f"!!! CAISR PIPELINE ERROR !!!")
        print(f"CAISR module crashed with error: {e}")
        print("Traceback:")
        traceback.print_exc()
        print("="*60 + "\n")
        caisr_results = {}

    _timer.stop()

    def _local_parse_caisr_stages(stage_csv_path: str, target_length: int) -> np.ndarray:
        if not stage_csv_path or not Path(stage_csv_path).exists():
            return np.zeros(target_length, dtype=int)
        try:
            df = pd.read_csv(stage_csv_path)
            caisr_to_std = {5.0: 0, 4.0: 4, 3.0: 3, 2.0: 2, 1.0: 1}
            # Конвертируем метки CAISR в стандартные (0,1,2,3,4)
            stages = df['stage'].map(caisr_to_std).fillna(0).astype(int).values
            
            
            stages_1hz = np.repeat(stages, 30)
            
            
            if len(stages_1hz) < target_length:
                stages_1hz = np.pad(stages_1hz, (0, target_length - len(stages_1hz)), 'constant', constant_values=0)
                
            return stages_1hz[:target_length]
        except Exception as e:
            print(f"Warning: Failed to parse CAISR stages: {e}")
            return np.zeros(target_length, dtype=int)

    sf = float(raw_export.info["sfreq"])
    raw_sec = int(raw_export.n_times // sf)
    common_sec = (raw_sec // 30) * 30

    yasa_1hz = np.repeat(hypno_int_30s, 30).astype(int)[:common_sec]
    caisr_1hz = _local_parse_caisr_stages(caisr_results.get("stages_csv", ""), common_sec)

    # SELECT PRIMARY HYPNOGRAM
    if primary_scoring.upper() == "CAISR":
        if caisr_results.get("stages_csv") and Path(caisr_results["stages_csv"]).exists():
            hypno_1hz_int = caisr_1hz
            if verbose: print(">>> Using CAISR output as the primary sleep hypnogram.")
        else:
            hypno_1hz_int = yasa_1hz
            if verbose: print(">>> CAISR stages not found or failed! Fallback to YASA output as primary.")
    else:
        hypno_1hz_int = yasa_1hz
        if verbose: print(">>> Using YASA output as the primary sleep hypnogram.")

    np.savetxt(hypno_txt, hypno_1hz_int, fmt="%d")

    hypno_int_30s_chosen = hypno_1hz_int[::30]
    hypno_str_30s_chosen = yasa.hypno_int_to_str(hypno_int_30s_chosen)
    counts = pd.Series(hypno_str_30s_chosen).value_counts()
    perc = counts / counts.sum() * 100

    caisr_nrem_events = None
    plm_metrics = None
    if caisr_results.get("limb_events_csv"):
        caisr_nrem_events = parse_caisr_nrem_events(caisr_results["limb_events_csv"], hypno_1hz_int)
        
        # Вычисляем индексы PLM
        plm_metrics = calculate_plm_metrics(caisr_nrem_events, hypno_1hz_int)
        if verbose:
            print(f">>> Calculated PLMI: {plm_metrics['PLMI']:.1f} events/hour (Total PLM: {plm_metrics['Total_PLM']})")

    sleep_stats = None
    stats_csv = None
    hypno_png = None
    sleep_report_png = None
    docx_report = None
    artifact_rejection = None
    final_pdf_path = None 
    
    rbdtector = None
    if run_rbdtector_headless:
        _timer.start("RBDtector headless analysis")
        rbd_hypno_input = hypno_1hz_int.copy()
        is_rem = (rbd_hypno_input == 4)
        
        if is_rem.any():
            edges = np.diff(np.concatenate(([0], is_rem, [0])))
            run_starts = np.where(edges == 1)[0]
            run_ends = np.where(edges == -1)[0]
            
            for start, end in zip(run_starts, run_ends):
                if (end - start) < 150:
                    rbd_hypno_input[start:end] = 0
            
            if verbose:
                print(">>> Cleaned tiny REM fragments from hypnogram for RBDtector analysis (marked as Wake).")
                
        if not (rbd_hypno_input == 4).any():
             if verbose:
                 print(">>> SKIPPING RBDtector: No consolidated REM blocks (>=150s) found to analyze.")
        else:
            if verbose:
                print(">>> Running RBDtector Headless Analysis...")
            try:
                rbdtector = _run_rbdtector_headless_case(
                    outdir=details_dir,
                    raw_for_edf=raw_export,
                    hypno_1hz_int=rbd_hypno_input, 
                    emg_label_in_edf=emg_alias,
                    verbose=verbose,
                )
            except Exception as e:
                print(f"RBDtector execution failed: {e}")
                traceback.print_exc()

    if run_rbdtector_headless:
        _timer.stop()

    spectral_pipe = None
    hspec_ok = False
    _timer.start("Spectrogram & visualizations")

    try:
        if verbose:
            print("Generating SpectralPipe hypnospectrogram with embedded RBD events...")

        save_hypnospectrogram_fallback(
            eeg_path=fif_path,
            outdir=details_dir,
            hypnotxt=hypno_txt,
            hypnofreq=hypno_freq,
            eeg_channel=eeg_alias_for_yasa,
            outpng=hspec_png,
            winsec=win_sec,
            freqrange=freq_range,
            cmap=cmap,
            overlap=overlap,
            signal_name=emg_alias,
            caisr_nrem_events=caisr_nrem_events
        )

        hspec_ok = Path(hspec_png).exists()

    except Exception as e:
        errtxt = details_dir / "spectralpipe_plot_error.txt"
        errtxt.write_text(traceback.format_exc(), encoding="utf-8")
        if verbose:
            print(f"SpectralPipe/fallback failed: {type(e).__name__}: {e}")

    fig1, ax1 = plt.subplots(figsize=(5, 5), dpi=300)
    ax1.pie(perc.values, labels=[f"{st} ({p:.1f}%)" for st, p in perc.items()], startangle=90)
    ax1.set_title(f"Stage composition (%) - {primary_scoring}")
    fig1.savefig(details_dir / "pie.png", bbox_inches="tight")
    plt.close(fig1)

    plot_sleep_stages(mff_path_primary.stem, hypno_report_stats(hypno_1hz_int), details_dir / "legacy_pie.png")

    if caisr_nrem_events is not None or (rbdtector and rbdtector.get("calculated_df") is not None):
        plot_combined_hypnogram(
            path=details_dir / f"{mff_path_primary.stem}_combined_hypnogram.png",
            hypno_1hz_int=hypno_1hz_int,
            start_dt=_choose_start_dt(raw),
            rbdtector_events=rbdtector["calculated_df"] if rbdtector else None,
            caisr_nrem_events=caisr_nrem_events,
            rbd_signal_name=emg_alias,
            title=f"Combined Analysis ({primary_scoring}): {mff_path_primary.stem}",
            verbose=verbose
        )

        try:
            plot_motor_events_standalone(
                out_path=caisr_motor_events_png,
                h=hypno_1hz_int,
                caisr_nrem_events=caisr_nrem_events,
                rbdtector_events=rbdtector["calculated_df"] if rbdtector else None,
                start_dt=_choose_start_dt(raw)
            )
        except Exception as e:
            if verbose:
                print(f"Failed to plot standalone motor events: {e}")

    plt.close("all")

    if hspec_ok and Path(details_dir / "hspec.png").exists():
        if verbose:
            print("Building combo from pie_png + hspec_png")
        im1 = Image.open(details_dir / "pie.png").convert("RGB")
        im2 = Image.open(details_dir / "hspec.png").convert("RGB")
        h_max = max(im1.height, im2.height)
        im1r = im1.resize((int(im1.width * h_max / im1.height), h_max), Image.Resampling.LANCZOS)
        im2r = im2.resize((int(im2.width * h_max / im2.height), h_max), Image.Resampling.LANCZOS)
        combined = Image.new("RGB", (im1r.width + im2r.width, h_max), (255, 255, 255))
        combined.paste(im1r, (0, 0))
        combined.paste(im2r, (im1r.width, 0))
        combined.save(outdir / "combo.png")
    else:
        if verbose:
            print("FALLBACK: copying pie_png -> combo_png")
        shutil.copy2(details_dir / "pie.png", outdir / "combo.png")

    preview_source = outdir / "combo.png" if show_preview else None

    _timer.stop()

    psd_stats_result = None
    if make_sleep_report:
        _timer.start("PSD stats & sleep report panel")
        psd_stats_result = _make_psd_stats_panel(
            raw=raw_export,
            hypno_1hz_int=hypno_1hz_int,
            subject_id=mff_path_primary.stem,
            outdir=outdir,
            eeg_channel_for_report=eeg_alias_for_report,
            separate_nrem_power_spectrum=True,
            use_yasa_art_detect=use_yasa_art_detect,
            yasa_art_epoch_sec=yasa_art_epoch_sec,
            yasa_art_window=yasa_art_window,
            yasa_art_method=yasa_art_method,
            yasa_art_threshold=yasa_art_threshold,
            yasa_art_include=yasa_art_include,
            verbose=verbose,
        )

        hypno_png = psd_stats_result["hypno_png"]
        stats_csv = psd_stats_result["stats_csv"]
        sleep_stats = psd_stats_result["sleep_stats"]
        artifact_rejection = psd_stats_result.get("artifact_rejection")

        generated_panel = Path(psd_stats_result["psd_stats_png"])
        if generated_panel != psd_stats_png:
            shutil.copy2(generated_panel, psd_stats_png)

        if hspec_ok and Path(hspec_png).exists():
            _stack_images_vertical(hspec_png, psd_stats_png, report_png)
        else:
            shutil.copy2(psd_stats_png, report_png)

        sleep_report_png = report_png
        if show_preview:
            preview_source = sleep_report_png

    if make_sleep_report:
        _timer.stop()

    docx_report_path = None
    if make_docx_report:
        _timer.start("DOCX/PDF report generation")
        if subject_age is None or sex is None:
            raise ValueError("subject_age and sex are required when make_docx_report=True")
        
        if hspec_ok and Path(hspec_png).exists():
            final_hypno_path = hspec_png
        elif run_rbdtector_headless and hypno_rbd_events_png.exists():
            final_hypno_path = hypno_rbd_events_png
        else:
            final_hypno_path = hypno_png

        if final_hypno_path is None or not Path(final_hypno_path).exists():
            final_hypno_path = outdir / f"{mff_path_primary.stem}_hypno.png"
            title_name = subject_name if subject_name else mff_path_primary.stem
            _save_hypnogram_png(final_hypno_path, hypno_1hz_int, title=title_name)


        docx_results = create_docx_sleep_report(
            report_type="detailed",
            subject_id=mff_path_primary.stem,
            subject_name=subject_name,
            outdir=outdir,
            hypno_1hz_int=hypno_1hz_int,
            subject_age=int(subject_age),
            sex=str(sex),
            hypno_png_path=Path(final_hypno_path),
            pie_png_path=legacy_pie_png,
            spo2_edf_path=spo2_edf_path,
            spo2_channel=spo2_channel,
            pleth_channel=pleth_channel,
            plm_metrics=plm_metrics,
            watchpat_pdf_path=append_pdf_path,
            motor_events_png_path=caisr_motor_events_png
        )
        
        # docx_results could be a tuple (en, he) or a single path if simple
        if not isinstance(docx_results, tuple):
            docx_results = (docx_results,)
            
        for d_path in docx_results:
            if not d_path or not Path(d_path).exists():
                continue
            docx_report_path = Path(d_path)
            try:
                from pypdf import PdfWriter
                base_pdf_path = docx_report_path.with_suffix(".pdf")
                if not _convert_docx_to_pdf_safe(docx_report_path, base_pdf_path, timeout_sec=120):
                    raise RuntimeError("docx2pdf conversion failed or timed out")

                if append_pdf_path and Path(append_pdf_path).exists():
                    merger = PdfWriter()
                    merger.append(str(base_pdf_path))
                    merger.append(str(append_pdf_path))
                    # determine language suffix
                    suffix = "_he" if "_he" in docx_report_path.stem else "_en" if "_en" in docx_report_path.stem else ""
                    final_pdf_path = outdir / f"sleep_report_{mff_path_primary.stem}_merged{suffix}.pdf"
                    merger.write(str(final_pdf_path))
                    merger.close()
                    if base_pdf_path.exists(): base_pdf_path.unlink()
            except Exception as e:
                print(f"Error during PDF conversion for {docx_report_path.name}: {e}")

    if show_preview and preview_source is not None:
        shutil.copy2(preview_source, preview_png)
    
    if make_docx_report:
        _timer.stop()

    # =========================================================================
    # PATIENT SUMMARY REPORT (always generated)
    # =========================================================================
    patient_report_png = None
    try:
        _timer.start("Patient summary report")
        from patient_report import generate_patient_report

        # --- pAHI from WatchPAT PDF ---
        _rdi = None
        _spo2_stats = None
        _wpat_hr_stats = None
        _pos_stats = None
        _resp_stats = None

        if append_pdf_path and Path(append_pdf_path).exists():
            try:
                _wpat = parse_watchpat_pdf(append_pdf_path)
                _rdi = float(_wpat["pAHI"]) if _wpat.get("pAHI") else None
                _pos_stats = _wpat.get("pos_stats")
                _resp_stats = _wpat.get("resp_stats")
                # SpO2 from WatchPAT is more reliable than EDF when WatchPAT is present
                if _wpat.get("mean_sat") and _wpat["mean_sat"] > 0:
                    _spo2_stats = {
                        "avg":       float(_wpat["mean_sat"]),
                        "min":       float(_wpat.get("min_sat", 0)),
                        "t_under90": float(_wpat.get("sat_below_90_pct", 0)),
                    }
                if _wpat.get("hr_stats"):
                    _wpat_hr_stats = _wpat["hr_stats"]
            except Exception as _we:
                print(f"[PatientReport] WatchPAT parse failed: {_we}")

        # Fallback: SpO2 from EDF if WatchPAT not available
        if _spo2_stats is None and make_docx_report and spo2_edf_path:
            try:
                avg_s, min_s, t90, _ = analyse_spo2(
                    spo2_edf_path,
                    spo2_channel=spo2_channel,
                    pleth_channel=pleth_channel,
                )
                _spo2_stats = {"avg": avg_s, "min": min_s, "t_under90": t90}
            except Exception:
                pass

        from cardio_metrics import compute_cardio_metrics, validate_and_merge_cv_stats
        _edf_cv_stats = compute_cardio_metrics(raw)
        _cv_stats = validate_and_merge_cv_stats(_wpat_hr_stats, _edf_cv_stats)

        patient_report_png = generate_patient_report(
            hypno_1hz_int=hypno_1hz_int,
            subject_id=mff_path_primary.stem,
            subject_name=subject_name,
            outdir=outdir,
            study_date=_choose_start_dt(raw).strftime("%d.%m.%Y") if _choose_start_dt(raw) else None,
            subject_age=int(subject_age) if subject_age is not None else None,
            sex=str(sex) if sex is not None else None,
            spo2_stats=_spo2_stats,
            plm_metrics=plm_metrics,
            rdi=_rdi,
            cv_stats=_cv_stats,
            pos_stats=_pos_stats,
            resp_stats=_resp_stats,
            caisr_nrem_events=caisr_nrem_events,
            rbdtector_events=rbdtector["calculated_df"] if "rbdtector" in locals() and rbdtector else None,
            start_dt=_choose_start_dt(raw),
        )
        _timer.stop()
    except Exception as _e:
        print(f"[Warning] Patient report generation failed: {_e}")
        traceback.print_exc()

    # =========================================================================
    # ГЕНЕРАЦИЯ АТЛАСА СОБЫТИЙ (PDF)
    # =========================================================================
    if make_sleep_report:
            _timer.start("Events Atlas PDF")
            events_pdf_path = outdir / f"{mff_path_primary.stem}_events_atlas.pdf"
            
            plm_l_alias = emg_alias
            plm_r_alias = None
            for ch in r.ch_names:
                ch_up = ch.upper()
                if 'LEFT_LEG' in ch_up or 'LAT' in ch_up:
                    plm_l_alias = ch
                elif 'RIGHT_LEG' in ch_up or 'RAT' in ch_up:
                    plm_r_alias = ch
                    
            try:
                print("\n[REPORT] Preparing filtered bipolar data for PDF Atlas...")
                raw_pdf = r.copy()
                raw_pdf.load_data()
                
                muscles = [ch for ch in [emg_alias, plm_l_alias, plm_r_alias] if ch and ch in raw_pdf.ch_names]
                if muscles:
                    raw_pdf.filter(l_freq=10.0, h_freq=None, picks=muscles, verbose=False)

                export_events_to_pdf(
                    raw=raw_pdf, 
                    out_path=events_pdf_path,
                    outdir=outdir,
                    caisr_nrem_events=caisr_nrem_events,
                    start_dt=_choose_start_dt(raw_pdf),
                    rbd_channel=emg_alias,            
                    rbd_match_label=emg_alias,      
                    plm_l_channel=plm_l_alias,          
                    plm_r_channel=plm_r_alias,
                    eeg_channels=[eeg_alias_F, eeg_alias_C, eeg_alias_O],
                    eog_channels=[eog1_alias, eog2_alias],
                    hypno_1hz_int=hypno_1hz_int, 
                    epoch_sec=30.0
                )
            except Exception as e:
                print(f"Error generating Events PDF Atlas: {e}")
                traceback.print_exc()
    
    if make_sleep_report:
            _timer.stop()

    _timer.report()

    summary = pd.DataFrame({"epochs": counts, "percent": perc}).loc[[s for s in ["W", "N1", "N2", "N3", "REM"] if s in counts.index]]
    cleanup_output_directory(outdir, verbose=verbose)
    # === CUMULATIVE DATABASE UPDATE ===
    if anchor_dir is not None:
        try:
            import datetime
            db_path = Path(anchor_dir) / "Cumulative_Sleep_Database.csv"
            
            # Gather metrics safely
            tst = sleep_stats.get('TST', 0) if sleep_stats else 0
            se = sleep_stats.get('SE', 0) if sleep_stats else 0
            sol = sleep_stats.get('SOL', 0) if sleep_stats else 0
            waso = sleep_stats.get('WASO', 0) if sleep_stats else 0
            
            n1_perc, n2_perc, n3_perc, rem_perc = 0, 0, 0, 0
            if summary is not None:
                for idx, row in summary.iterrows():
                    if idx == 'N1': n1_perc = row.get('percent', 0)
                    elif idx == 'N2': n2_perc = row.get('percent', 0)
                    elif idx == 'N3': n3_perc = row.get('percent', 0)
                    elif idx == 'REM': rem_perc = row.get('percent', 0)
            
            plmi = plm_metrics.get('plm_index', 0) if plm_metrics else 0
            
            rbd_idx = 0
            try:
                if "rbdtector" in locals() and rbdtector is not None and "calculated_df" in rbdtector:
                    df_rbd = rbdtector["calculated_df"]
                    if df_rbd is not None and not df_rbd.empty and "RBD Index" in df_rbd.columns:
                        rbd_idx = float(df_rbd["RBD Index"].iloc[0])
            except: pass
            
            age_val = subject_age if subject_age is not None else 0
            sex_val = sex if sex is not None else "Unknown"
            
            _loc_rdi = locals().get("_rdi")
            _loc_spo2 = locals().get("_spo2_stats") or {}
            
            pahi_val = _loc_rdi if _loc_rdi is not None else 0.0
            spo2_avg = _loc_spo2.get('avg', 0.0)
            spo2_min = _loc_spo2.get('min', 0.0)
            spo2_t90 = _loc_spo2.get('t_under90', 0.0)

            psd_metrics = {}
            
            # --- Advanced EEG Features (Spindles, Slow Waves) ---
            try:
                if "raw" in locals() and raw is not None and "hypno_1hz_int" in locals() and hypno_1hz_int is not None:
                    # Убедимся, что raw загружен в память (preload=True) для корректного get_data()
                    if not raw.preload:
                        raw.load_data()
                    
                    from advanced_eeg_pipeline import run_advanced_eeg_analysis
                    sfreq = raw.info["sfreq"]
                    adv_metrics = run_advanced_eeg_analysis(raw, hypno_1hz_int, sfreq)
                    psd_metrics.update(adv_metrics)
            except Exception as adv_e:
                print(f"[!] Failed to compute Advanced EEG metrics: {adv_e}")
            # ----------------------------------------------------

            if "epochs" in locals() and epochs is not None:
                try:
                    bands = [
                        (0.5, 4, 'Delta'),
                        (4, 8, 'Theta'),
                        (8, 12, 'Alpha'),
                        (12, 16, 'Sigma'),
                        (16, 30, 'Beta')
                    ]
                    for st_name, st_val in [("N2", 2), ("N3", 3)]:
                        if len(epochs[st_name]) > 0:
                            data_st = epochs[st_name].get_data(picks=eeg_channel_for_report, copy=True) * 1e6 # Convert Volts to microVolts
                            sfreq = epochs.info["sfreq"]
                            bp_abs = yasa.bandpower(data_st, sfreq=sfreq, ch_names=[eeg_channel_for_report], bands=bands, relative=False)
                            bp_rel = yasa.bandpower(data_st, sfreq=sfreq, ch_names=[eeg_channel_for_report], bands=bands, relative=True)
                            
                            bp_abs_mean = bp_abs.mean(numeric_only=True)
                            bp_rel_mean = bp_rel.mean(numeric_only=True)
                            
                            for b_name in ['Delta', 'Theta', 'Alpha', 'Sigma', 'Beta']:
                                psd_metrics[f"{st_name}_{b_name}_abs_uV2"] = bp_abs_mean.get(b_name, 0.0)
                                psd_metrics[f"{st_name}_{b_name}_rel_perc"] = bp_rel_mean.get(b_name, 0.0) * 100.0
                except Exception as e:
                    print(f"[!] Failed to compute PSD for database: {e}")
            
            import socket
            
            new_row = {
                "Subject_ID": subject_id,
                "Computer_Name": socket.gethostname(),
                "Date_Processed": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "Scoring_System": primary_scoring,
                "Age": age_val,
                "Sex": sex_val,
                "TST_min": tst,
                "SE_perc": se,
                "SOL_min": sol,
                "WASO_min": waso,
                "N1_perc": n1_perc,
                "N2_perc": n2_perc,
                "N3_perc": n3_perc,
                "REM_perc": rem_perc,
                "PLMI": plmi,
                "pAHI": pahi_val,
                "SpO2_Avg": spo2_avg,
                "SpO2_Min": spo2_min,
                "SpO2_T90": spo2_t90,
                "RBD_Index": rbd_idx
            }
            new_row.update(psd_metrics)
            
            new_df = pd.DataFrame([new_row])
            if db_path.exists():
                db_df = pd.read_csv(db_path)
                db_df = pd.concat([db_df, new_df], ignore_index=True)
            else:
                db_df = new_df
                
            db_df.to_csv(db_path, index=False)
            print(f">>> Cumulative database updated at {db_path.name}")
            
            try:
                from gdrive_sync import sync_csv_to_gdrive
                sync_csv_to_gdrive(str(db_path), "1LHiTuI07iXg7I7AFZ-MMtV3UKNWkkINM")
            except Exception as ge:
                print(f"[!] GDrive Sync Error: {ge}")
                
        except Exception as e:
            print(f"[!] Failed to update cumulative database: {e}")
    # ==================================

    return {
        "paths": {
            "outdir": outdir,
            "hypno_txt": hypno_txt,
            "pie_png": pie_png,
            "legacy_sleep_stages_png": legacy_pie_png,
            "hypnospec_png": hspec_png,
            "combo_png": combo_png,
            "fif_path": fif_path,
            "hypno_png": hypno_png,
            "hypno_rbd_png": hypno_rbd_events_png if hypno_rbd_events_png.exists() else None,
            "stats_csv": stats_csv,
            "sleep_report_png": sleep_report_png,
            "psd_stats_png": psd_stats_png if make_sleep_report else None,
            "artifact_csv": artifact_csv if (make_sleep_report and use_yasa_art_detect) else None,
            "preview_png": preview_png if show_preview else None,
            "docx_report": docx_report_path,
            "pdf_report": final_pdf_path,
            "patient_report_png": patient_report_png,
        },
        "cfg_used": {
            "side": side,
            "site": site,
            "need_loaded": need,
            "export_order": export_order,
            "report_channel": eeg_alias_for_report,
            "apply_prefilter": apply_prefilter,
            "prefilter_l_freq": prefilter_l_freq,
            "prefilter_h_freq": prefilter_h_freq,
            "notch_base_freq": notch_base_freq,
            "use_yasa_art_detect": use_yasa_art_detect,
            "yasa_art_epoch_sec": yasa_art_epoch_sec,
            "yasa_art_window": yasa_art_window,
            "yasa_art_method": yasa_art_method,
            "yasa_art_threshold": yasa_art_threshold,
            "yasa_art_include": yasa_art_include,
            "montage_type": montage_type,
            "primary_scoring": primary_scoring
        },
        "raw_min_loaded": raw,
        "raw_export": raw_export,
        "raw_yasa": raw_yasa,
        "hyp": hyp,
        "hypno_int_30s": hypno_int_30s,
        "hypno_1hz_int": hypno_1hz_int,
        "perc": perc,
        "summary": summary,
        "sleep_stats": sleep_stats,
        "plm_metrics": plm_metrics,
        "artifact_rejection": artifact_rejection,
        "legacy_hypno_stats": hypno_report_stats(hypno_1hz_int),
        "caisr_results": caisr_results, 
    }