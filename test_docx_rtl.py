import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = docx.Document()
p = doc.add_paragraph()
if True: # lang == 'he'
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("שלום עולם")
run.font.rtl = True

table = doc.add_table(rows=2, cols=2)
tblPr = table._element.xpath('w:tblPr')
if tblPr:
    bidi = OxmlElement('w:bidiVisual')
    tblPr[0].append(bidi)

cell = table.cell(0,0)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("שורה 1 עמודה 1")
run.font.rtl = True

doc.save("test_rtl.docx")
print("Saved test_rtl.docx")
